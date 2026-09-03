"""Build meeting_prep_short.docx -- a short, focused replacement for the
longer meeting_prep_v4.docx.

No cheat-sheet page, no full "since last meeting" walkthrough. Four
sections, in order:
  1. Indiana Payment Order (answering Hall's question) -- content pulled
     from PROJECT_CONTEXT.md's 2026-09-02 Task 1b section, with the
     canal/robustness chart embedded.
  2. Why Didn't the City Itself Default? -- content pulled from
     PROJECT_CONTEXT.md's 2026-09-02 Task 2 section, with the
     Philadelphia-vs-Pennsylvania and Pittsburgh-vs-Pennsylvania charts
     embedded.
  3. What the Research Paper Will Cover -- a section-by-section walkthrough
     of the proposed paper (Introduction, the problem, data/method, main
     result, city result, robustness, discussion, conclusion), each as a
     short description plus bullet points of what goes in it.
  4. How I'll Run the Meeting -- a spoken-flow guide (first-person,
     casual, 2-3 sentence bullets), NOT a formal numbered outline.

Visual style (accent color, SAY THIS boxes as bordered paragraphs,
muted-gray italics for private asides) is carried over from
build_meeting_prep_v4.py unchanged. Boxes are bordered paragraphs, not
tables (a table-based box silently clipped its last line under the
docx->PDF renderer -- see prior passes' notes).
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "output"
CHART_DIR = OUTPUT_DIR / "charts_v3_cropped"

INK = RGBColor(0x0B, 0x0B, 0x0B)
MUTED = RGBColor(0x52, 0x51, 0x4E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

ACCENT = "1B4B66"
ACCENT_RGB = RGBColor(0x1B, 0x4B, 0x66)
ACCENT_TINT = "E9F1F5"

embedded_charts = []
missing_charts = []


# ---------------------------------------------------------------------------
# Helpers (carried over from build_meeting_prep_v4.py)
# ---------------------------------------------------------------------------

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(4)


def muted_note_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(4)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_chart(doc, filename, caption, width_in=6.0):
    path = CHART_DIR / filename
    if path.exists():
        pic_p = doc.add_paragraph()
        pic_p.paragraph_format.space_before = Pt(4)
        pic_p.paragraph_format.space_after = Pt(0)
        run = pic_p.add_run()
        run.add_picture(str(path), width=Inches(width_in))
        embedded_charts.append(str(path))
        add_caption(doc, caption)
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[chart not found - check output/charts_v3_cropped/: {filename}]")
        run.bold = True
        run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
        missing_charts.append(filename)


def set_style_bottom_rule(style, color_hex, size="6"):
    pPr = style.element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_numbers(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


# --- Bordered-paragraph box (not a table) ---

def start_box(doc):
    return len(doc.paragraphs)


def end_box(doc, start_index, fill_hex, border_hex, indent_pt=14, label_fill=None,
            label_text_white=False):
    paras = doc.paragraphs[start_index:]
    n = len(paras)
    for i, p in enumerate(paras):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        edge_specs = [("left", "16"), ("right", "16")]
        if i == 0:
            edge_specs.append(("top", "16"))
        if i == n - 1:
            edge_specs.append(("bottom", "16"))
        for edge, sz in edge_specs:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), sz)
            el.set(qn("w:space"), "8")
            el.set(qn("w:color"), border_hex)
            pBdr.append(el)
        pPr.append(pBdr)
        this_fill = label_fill if (label_fill and i == 0) else fill_hex
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), this_fill)
        pPr.append(shd)
        p.paragraph_format.left_indent = Pt(indent_pt)
        p.paragraph_format.right_indent = Pt(indent_pt)
        p.paragraph_format.space_before = Pt(6 if i > 0 else 7)
        p.paragraph_format.space_after = Pt(6 if i < n - 1 else 7)
        if label_fill and i == 0 and label_text_white:
            for run in p.runs:
                run.font.color.rgb = WHITE


def _keep_para_together(p, keep_next):
    """Prevent this paragraph from splitting across a page break, and
    (optionally) bind it to the next one -- so a SAY THIS box moves whole
    to the next page rather than breaking across the boundary."""
    pf = p.paragraph_format
    pf.keep_together = True
    if keep_next:
        pf.keep_with_next = True


def say_this(doc, text):
    start = start_box(doc)
    label = doc.add_paragraph()
    run = label.add_run("SAY THIS")
    run.bold = True
    run.font.small_caps = True
    run.font.size = Pt(10)
    body = doc.add_paragraph()
    run2 = body.add_run(text)
    run2.italic = True
    end_box(doc, start, ACCENT_TINT, ACCENT, label_fill=ACCENT, label_text_white=True)
    box_paras = doc.paragraphs[start:]
    for i, bp in enumerate(box_paras):
        _keep_para_together(bp, keep_next=(i < len(box_paras) - 1))
    doc.paragraphs[-1].paragraph_format.space_after = Pt(12)


def paper_part(doc, heading, sentences, bullets):
    """One part of the proposed paper: a Heading-2 sub-heading, a
    one-or-two-sentence description, then bullet points of what goes in
    it. Sub-heading kept with the description so it doesn't strand at a
    page break."""
    h = doc.add_heading(heading, level=2)
    h.paragraph_format.keep_with_next = True
    d = doc.add_paragraph()
    d.paragraph_format.keep_with_next = True
    d.paragraph_format.space_after = Pt(3)
    d.add_run(sentences)
    add_bullets(doc, bullets)


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)

    h1 = doc.styles["Heading 1"]
    h1.font.color.rgb = ACCENT_RGB
    h1.font.size = Pt(17)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    set_style_bottom_rule(h1, ACCENT, size="6")

    h2 = doc.styles["Heading 2"]
    h2.font.color.rgb = ACCENT_RGB
    h2.font.size = Pt(12.5)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(3)

    title = doc.add_heading("1840s State Debt Credibility — Short Update", level=0)
    title.paragraph_format.space_after = Pt(4)
    lead = doc.add_paragraph()
    r = lead.add_run(
        "Four things: the Indiana payment-order answer to Hall's question, why "
        "Philadelphia itself never defaulted, how the paper will flow, and how I'll run "
        "the meeting."
    )
    r.italic = True
    r.font.color.rgb = MUTED

    # ------------------------------------------------------------------
    # Section 1 -- Indiana Payment Order
    # ------------------------------------------------------------------
    doc.add_heading("1. The Indiana Payment Order — Answering Hall's Question", level=1)
    doc.add_paragraph(
        "Last meeting, Hall asked which of Indiana's canal bonds actually got paid first. "
        "Here's the clean version."
    )
    add_bullets(doc, [
        "The line-up: two bonds got paid first — S-0490 and S-0500, the ones labeled "
        "\"preferred.\" Two got paid only after those — S-0480 and S-0506, the \"deferred\" "
        "ones. A fifth bond, S-0470, labeled just \"Indiana Canal,\" doesn't fit either group: "
        "it traded far higher than all four and looks like a separate, better-secured "
        "instrument, probably the part of the debt the state itself still stood behind.",
        "There's even a pecking order inside \"preferred\": the plain preferred bond (S-0490) "
        "traded higher than the one labeled \"special preferred\" (S-0500) — roughly $42 vs. "
        "$20 median price. So the market wasn't just splitting these into two tiers, it was "
        "pricing finer distinctions than that.",
        "What we can't tell: whether a shortfall got split strictly (the preferred group paid "
        "in full before the deferred group sees a cent) or proportionally (both groups take a "
        "haircut at the same time). No source we found states the rule directly.",
        "What the prices themselves imply: the deferred bonds traded at a real, positive price "
        "— not near zero — so the market expected those holders to get something eventually. "
        "And even the preferred bonds only traded around 20 to 50 cents on the dollar, well "
        "below full value. So both groups were priced as taking a loss — not a clean \"the "
        "senior bond gets everything, the junior bond gets wiped out\" split.",
        "Issuance timing: we couldn't pin it down. The codebook has no issue-date field, and "
        "the legislative histories don't give a rollout schedule. The five bonds first show up "
        "on the exchange at staggered dates in 1850–51 — but that's three years after the 1847 "
        "restructuring act, and reflects when they started trading, not when they were issued. "
        "This is a genuine gap in the historical record, not something more digging is likely "
        "to fix.",
    ])
    add_chart(
        doc, "chart_canal_robustness.png",
        "Right panel: Indiana's preferred tranches (green) vs. deferred tranches (orange) of "
        "the same 1847 restructured canal loan — a ~40-percentage-point yield gap, the largest "
        "in the project. Left panel: New York's canal bonds, for comparison.",
    )
    say_this(
        doc,
        "So last time you asked which Indiana bonds got paid first — I've got a clear answer "
        "now. Two of them were \"preferred\" and got paid first; two were \"deferred\" and got "
        "paid after; and there's a fifth one that doesn't really fit either bucket and looks "
        "like a separate, safer instrument. There's even a sub-ranking inside the preferred "
        "group, so the market was pricing fine distinctions. What I can't tell you is the exact "
        "rule for how a shortfall gets split between them — nothing in the historical record "
        "spells that out — but the price data itself is telling: both the preferred and the "
        "deferred bonds traded well below full value, which suggests the market expected both "
        "sides to take some kind of hit, not a clean \"senior gets everything, junior gets "
        "nothing\" split. And I couldn't pin down exactly when the bonds were issued — the "
        "records just don't have it.",
    )

    # ------------------------------------------------------------------
    # Section 2 -- Why Didn't the City Itself Default?
    # ------------------------------------------------------------------
    doc.add_heading("2. Why Didn't the City Itself Default?", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "This is a different question from the market-pricing finding. That earlier work showed "
        "investors didn't punish Philadelphia for Pennsylvania's default. This explains why "
        "Philadelphia never got into trouble in the first place — why the city itself didn't "
        "default alongside the state. "
    )
    p.add_run(
        "Keep the two separate: \"investors didn't punish the city\" is the price finding; "
        "\"the city stayed solvent because it was a legally separate entity\" is this."
    ).italic = True
    add_bullets(doc, [
        "Philadelphia was legally its own thing. A city in this era was a separate corporation "
        "from the state — its own budget, its own power to tax, borrow, and spend, not a branch "
        "of the state government.",
        "The state's debt paid for state projects. Pennsylvania borrowed to build the Main Line "
        "of Public Works — a state-owned canal-and-railroad route from Philadelphia to "
        "Pittsburgh — plus other state canals and investments in banks. Philadelphia the city "
        "was never a co-obligor or guarantor on any of it; the Main Line was a state asset.",
        "The city's debt paid for city things. Water works, gas works, wharves, public "
        "buildings — paid for out of the city's own property taxes and water and gas fees. That "
        "revenue was local and held up fine; it didn't depend on the state canal tolls and "
        "state taxes that collapsed.",
        "Different entity, different money — so the city never got pulled into the state's mess.",
    ])
    muted_note_bullet(
        doc,
        "Note to self: no single old source states this contrast in one sentence — it's pieced "
        "together from how the finances were structured plus our own price data. And I only "
        "checked the legal setup for Philadelphia specifically; Pittsburgh's price pattern "
        "matches but I didn't separately chase down its legal structure.",
    )
    add_chart(
        doc, "chart_city_vs_state.png",
        "Philadelphia (city, purple) barely reacts to Pennsylvania's default while the state's "
        "own yield (blue) spikes toward 20%.",
    )
    add_chart(
        doc, "chart_pittsburgh_vs_pa.png",
        "Pittsburgh (city, green) vs. Pennsylvania (state, blue) — the same ~2.65-point gap as "
        "Philadelphia, a second Pennsylvania city showing the same pattern.",
    )
    say_this(
        doc,
        "One more thing worth adding here — why didn't the city default, not just why didn't "
        "the market punish it. Philadelphia was legally its own thing, separate from the state, "
        "with its own taxing power and its own budget. The state's debt paid for state canals "
        "and railroads that Philadelphia had nothing to do with; the city's debt paid for city "
        "things, funded by the city's own taxes. Different entity, different money, so it never "
        "actually got pulled into the state's mess in the first place. I'll flag this is "
        "inference from how the finances were structured, not a single source that spells it "
        "out directly — but it's a solid, consistent picture.",
    )

    # ------------------------------------------------------------------
    # Section 3 -- What the Research Paper Will Cover
    # ------------------------------------------------------------------
    doc.add_heading("3. What the Research Paper Will Cover — Section by Section", level=1)
    lead3 = doc.add_paragraph()
    lr = lead3.add_run(
        "How the paper flows, part by part — a short description of each section and bullet "
        "points for what goes in it. Order is a proposal, not locked."
    )
    lr.italic = True
    lr.font.color.rgb = MUTED

    paper_part(
        doc,
        "Introduction",
        "Open on why this matters right now, then the historical parallel, then the question "
        "and a preview of the answer.",
        [
            "The modern hook: the EU is currently debating \"blue bonds\" — joint EU borrowing "
            "— and the core dispute is whether a credible \"no bailout\" rule actually forces "
            "member states to keep their own finances in order (Hanno Lustig, March 2026).",
            "The historical parallel: the United States faced the same choice in the early "
            "1840s, when nine states plus the Florida Territory defaulted and the federal "
            "government refused to assume their debts.",
            "The research question: did bond investors actually believe that refusal was "
            "permanent? If they did, the states that defaulted should have kept paying higher "
            "interest rates for years afterward.",
            "Preview of the answer: yes — and the market turned out to be a more careful "
            "judge than a simple panic story would predict.",
        ],
    )
    paper_part(
        doc,
        "The Problem — What Went Wrong in the 1840s",
        "The setup the rest of the paper analyzes: how the states got into debt, how the "
        "crisis hit, and why the federal government's response is the interesting part.",
        [
            "The 1830s borrowing boom: states borrowed heavily from British and Dutch "
            "investors to build canals, railroads, and banks.",
            "The Panic of 1837 crushed state tax revenue and trade.",
            "1841–1843: nine states plus the Florida Territory defaulted or suspended interest "
            "payments.",
            "The federal government refused to take on the debts — a deliberate break from the "
            "1790 precedent, when Hamilton had the federal government assume the states' "
            "Revolutionary War debt.",
            "Why this is the question worth studying: a \"no bailout\" rule only disciplines "
            "borrowers if lenders actually believe it will hold.",
        ],
    )
    paper_part(
        doc,
        "Data and Method",
        "What the evidence is and how we turn it into a test.",
        [
            "Source: transatlantic bond price records from the New York and Philadelphia "
            "exchanges, 1830s–1850s (EH.net Early U.S. Securities Prices).",
            "We compute interest rates (yields) ourselves, since the records give prices, not "
            "rates.",
            "Three groups of states: defaulted, risky-but-survived, and never-really-at-risk "
            "(the safe baseline).",
            "Two separate before/after moments, kept distinct so they don't get confused: the "
            "1837 crash itself, and the 1843 federal no-bailout signal.",
        ],
    )
    paper_part(
        doc,
        "Main Result — Did the No-Bailout Signal Show Up in State Interest Rates?",
        "The core finding: comparing the three groups of states after the federal refusal.",
        [
            "After the signal, the states that defaulted (Pennsylvania, Indiana) kept paying a "
            "premium over the safe baseline (Ohio) — the market did not forgive and forget.",
            "The premium is real but modest for Pennsylvania (~2 percentage points) once a "
            "yield-calculation artifact is corrected; it is larger and more persistent for "
            "Indiana.",
            "The \"risky but survived\" states land in between — with honest caveats (New "
            "York's usable data starts late; Alabama's classification is still being "
            "confirmed).",
            "Read: the market treated the federal refusal as credible.",
        ],
    )
    paper_part(
        doc,
        "City Result — Does a State's Default Spill Onto Its Own City?",
        "A distinct question the same data can answer: whether the punishment spread from a "
        "state to the city inside it.",
        [
            "Philadelphia's borrowing cost barely moved when Pennsylvania defaulted, while the "
            "state's own rate spiked — no spillover.",
            "Pittsburgh replicates it; Cincinnati and New York City (whose states were fine) "
            "show no gap either — the city/state gap only opens when the state itself is in "
            "trouble.",
            "The mechanical reason: the city was a legally separate entity with its own "
            "revenue, not on the hook for the state's debt.",
            "Read: the market told state credit and city credit apart, even for the same "
            "place.",
        ],
    )
    paper_part(
        doc,
        "Robustness — Canal Bonds and the Indiana Tranches",
        "An independent check on the main result, using a different kind of state debt.",
        [
            "Bonds backed specifically by canal toll revenue, analyzed separately from "
            "general state debt.",
            "Indiana's 1847 restructuring split its canal debt into \"preferred\" and "
            "\"deferred\" slices; the market priced the difference by about 40 percentage "
            "points — the largest gap anywhere in the project.",
            "Shows the market pricing fine seniority distinctions, not reacting bluntly.",
        ],
    )
    paper_part(
        doc,
        "Discussion",
        "What the pattern across all the results means, tied back to the modern debate.",
        [
            "The through-line: the market was a careful discriminator — it separated the panic "
            "from the policy, the state from its city, and senior debt from junior debt.",
            "Back to the EU: this is what a credible no-bailout commitment looks like when "
            "investors believe it — sustained higher costs for the specific borrowers who "
            "defaulted, not blind contagion to everyone nearby.",
        ],
    )
    paper_part(
        doc,
        "Conclusion and Limitations",
        "Restate the finding and be upfront about what the data can't do.",
        [
            "The 1840s bond market treated the federal no-bailout stance as credible, and "
            "priced credit distinctions finely.",
            "Limitations stated plainly: some bonds have coverage gaps in the key windows; the "
            "exact anchor date rests partly on secondary sources; the Alabama classification "
            "and the Indiana deferred-coupon question are still open.",
        ],
    )

    # ------------------------------------------------------------------
    # Section 4 -- How I'll Run the Meeting
    # ------------------------------------------------------------------
    doc.add_heading("4. How I'll Run the Meeting", level=1)
    intro = doc.add_paragraph()
    ir = intro.add_run(
        "Spoken-flow notes — how I'll actually open and steer the conversation. Not a formal "
        "outline; this is what I'll say."
    )
    ir.italic = True
    ir.font.color.rgb = MUTED
    add_bullets(doc, [
        "Opening: \"It's been a while, so I kept this short on purpose — I'm not going to walk "
        "through everything. Really just two answers I owe you, plus how I'm thinking about the "
        "paper.\" Set the expectation that this is a focused check-in, not a status marathon.",
        "Into the Indiana finding: \"Last time, Hall, you asked which Indiana bonds got paid "
        "first. I can map that out clearly now.\" Give the preferred/deferred split and the "
        "odd-one-out fifth bond, then hand it back: \"Does that answer what you were getting "
        "at, or were you after something more specific?\" Flag the one thing I still can't "
        "resolve — the exact shortfall-split rule — and note the price data leans toward both "
        "sides taking a hit.",
        "Into the city-default question: \"The other thing is why Philadelphia itself never "
        "defaulted — not why the market didn't punish it, which we already showed, but why the "
        "city never actually got into trouble.\" Walk the separate-corporation point, the "
        "state-debt-vs-city-debt split, and be upfront that it's inference from structure, not "
        "a single quotable source. Show the two charts as the supporting picture.",
        "Pivot to the paper: \"If you're good with where things are, here's roughly how I see "
        "it coming together\" — then walk the section-by-section flow from Section 3 above. "
        "Keep it loose: \"Nothing's locked; I mainly want your read on the order and the "
        "framing.\"",
        "Closing: \"So the two things I actually need are the Alabama call and your read on "
        "the Indiana deferred-coupon question — everything else is ready to write.\" I'll send "
        "a short recap of whatever we decide and start drafting the introduction next.",
    ])

    add_page_numbers(doc)

    out_path = OUTPUT_DIR / "meeting_prep_short.docx"
    doc.save(str(out_path))
    print(f"saved -> {out_path}")
    print("\nEmbedded charts:")
    for c in embedded_charts:
        print(f"  - {c}")
    if missing_charts:
        print("\nMISSING charts (placeholder inserted):")
        for c in missing_charts:
            print(f"  - {c}")


if __name__ == "__main__":
    main()
