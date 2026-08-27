"""Build meeting_prep_v2.docx -- plain-language rewrite + spoken script.

v2 of the meeting-prep talking doc (2026-08-27). Keeps meeting_prep_final.docx
(v1) untouched as a separate deliverable. Two changes from v1:

1. Plain-language rewrite of every section (jargon explained inline the
   first time it's used), plus a new "Background" section up front, written
   for Amey to read privately before the meeting -- not to be presented.
2. A "SAY THIS:" spoken-script block after every presented section, written
   as natural first-person speech, not slide bullets.

All numbers pulled fresh from PROJECT_CONTEXT.md (confirmed unchanged since
the last full read via `git log`/`git status` before this script was
written), not from the v1 docx.

Boxes (Background, SAY THIS, the agenda box, and the two-decision box) are
built from bordered paragraphs, not a single-cell table -- an earlier pass
found that a table-based box silently clipped its last line under Pages'
docx renderer even with `w:cantSplit` set (confirmed via raw XML inspection
that the content existed but wasn't rendered). Paragraph borders don't have
that failure mode.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "output"

INK = RGBColor(0x0B, 0x0B, 0x0B)
MUTED = RGBColor(0x52, 0x51, 0x4E)

BG_FILL = "EDE8DC"       # Background section: warm tan, visually distinct
BG_BORDER = "8C7B57"
SAY_FILL = "E8F0FE"      # SAY THIS blocks: light blue, consistent throughout
SAY_BORDER = "3B6EA5"
AGENDA_FILL = "EEF2F7"
AGENDA_BORDER = "6B7A8F"
BOX_FILL = "FFF7DB"      # Two-decision box: highlighted yellow
BOX_BORDER = "B8860B"

embedded_charts = []
missing_charts = []


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_chart(doc, filename, caption):
    path = OUTPUT_DIR / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.3))
        embedded_charts.append(str(path))
        add_caption(doc, caption)
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[chart not found — check output/: {filename}]")
        run.bold = True
        run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
        missing_charts.append(filename)


def add_toc(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose \"Update Field\" to generate the table of contents."
    fldChar_sep.append(placeholder)
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fldChar_begin)
    r.append(instrText)
    r.append(fldChar_sep)
    r.append(fldChar_end)


# --- Bordered-paragraph box (see module docstring for why not a table) ---

def start_box(doc):
    return len(doc.paragraphs)


def end_box(doc, start_index, fill_hex, border_hex, indent_pt=14):
    paras = doc.paragraphs[start_index:]
    n = len(paras)
    for i, p in enumerate(paras):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        edges = ["left", "right"]
        if i == 0:
            edges.append("top")
        if i == n - 1:
            edges.append("bottom")
        for edge in edges:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "18")
            el.set(qn("w:space"), "8")
            el.set(qn("w:color"), border_hex)
            pBdr.append(el)
        pPr.append(pBdr)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        pPr.append(shd)
        p.paragraph_format.left_indent = Pt(indent_pt)
        p.paragraph_format.right_indent = Pt(indent_pt)
        p.paragraph_format.space_before = Pt(8 if i == 0 else 4)
        p.paragraph_format.space_after = Pt(8 if i == n - 1 else 4)


def box_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return p


def box_bullet(doc, text):
    p = doc.add_paragraph()
    p.add_run("•  " + text)
    return p


def box_paragraph(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    return p


def say_this(doc, text):
    """Consistent spoken-script block. Same style every time so it's
    findable at a glance while presenting live."""
    start = start_box(doc)
    label = doc.add_paragraph()
    run = label.add_run("SAY THIS:")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x3D, 0x5C)
    body = doc.add_paragraph()
    run2 = body.add_run(text)
    run2.italic = True
    end_box(doc, start, SAY_FILL, SAY_BORDER)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = INK

    # ------------------------------------------------------------------
    # Title
    # ------------------------------------------------------------------
    doc.add_heading("1840s State Debt Credibility — Progress Since Last Meeting", level=0)

    # ------------------------------------------------------------------
    # Section 0: Background (private -- not spoken)
    # ------------------------------------------------------------------
    doc.add_heading("Background — What Is This Project, Actually", level=1)
    bg_start = start_box(doc)
    box_paragraph(
        doc,
        "(Read this before the meeting — it's for you, not something to say out loud.)",
        italic=True, bold=True,
    )
    box_paragraph(
        doc,
        "In the 1830s and early 1840s, American states borrowed huge amounts of money from "
        "British and Dutch investors to pay for canals, railroads, and state-owned banks. Then a "
        "financial crash hit in 1837, state tax revenue collapsed, and by the early 1840s nine "
        "states plus the Florida Territory couldn't make their bond payments — they defaulted. "
        "Congress and President Tyler's administration considered bailing them out, the way the "
        "federal government had once bailed out states before, back in 1790. This time, the "
        "federal government refused."
    )
    box_paragraph(
        doc,
        "This project asks a simple question: did investors actually believe the federal "
        "government meant it? A promise not to bail someone out only works if people believe it. "
        "If bond buyers in the 1840s believed the \"no bailout\" promise was real, states that "
        "defaulted should have kept paying higher interest rates for years afterward — investors "
        "would keep treating them as riskier, since no rescue was coming. If investors didn't "
        "really believe it, defaulted states' rates should have bounced back to normal fairly "
        "quickly."
    )
    box_paragraph(
        doc,
        "Why this old story matters right now: Stanford economist Hanno Lustig wrote a blog post "
        "in March 2026, \"Blue Bonds for Europe,\" arguing the EU is having almost the exact same "
        "argument today — should Europe issue joint debt with an implicit promise that struggling "
        "countries will get bailed out, or does fiscal discipline only work if the \"no bailout\" "
        "promise is credible? This project is basically a real historical test of that argument, "
        "using 200-year-old data instead of a hypothetical."
    )
    box_paragraph(
        doc,
        "What \"the data\" actually is, in the simplest terms: old bond price records — weekly "
        "price quotes for specific state bonds, pulled from newspapers and stock exchanges in "
        "Philadelphia and New York in the 1830s-1850s, since digitized into a database (EH.net). A "
        "bond's price moving up or down tells you how risky investors thought it was at that "
        "moment. We convert those prices into interest rates (yields) so they're comparable across "
        "different bonds and states."
    )
    box_paragraph(
        doc,
        "We sorted states into three groups. First, states that actually defaulted and stopped "
        "paying — Pennsylvania and Indiana for sure, and likely Alabama too (still to be "
        "confirmed — see the decision box later in this document). Second, a state that got "
        "scared and looked shaky but paid its bondholders the whole time — New York. Third, a "
        "state that was never really at serious risk — Ohio, which we use as the safe baseline "
        "everyone else gets compared against."
    )
    box_paragraph(
        doc,
        "Partway through, Hall raised a sharper question: does the punishment for a state's "
        "default also land on the CITY inside that state — like Philadelphia, sitting inside a "
        "defaulted Pennsylvania — even though the city government itself did nothing wrong and "
        "kept paying its own debts on time? That became a second, related track of this project, "
        "alongside the original state-level question."
    )
    box_heading(doc, "The short version of what we found:")
    box_bullet(
        doc,
        "States that defaulted kept paying more in interest for years afterward — the market "
        "didn't forgive and forget.",
    )
    box_bullet(
        doc,
        "The cities inside those states did NOT get punished — investors could tell the "
        "difference between a state's credit and its own city's credit, even sharing a name and a "
        "place.",
    )
    box_bullet(
        doc,
        "One huge extra finding worth flagging: in Indiana, one slice of a restructured loan got "
        "paid before another slice (a \"tranche\" just means one slice of a loan with its own "
        "place in line) — and the market priced that difference by more than anything else we "
        "found in this whole project.",
    )
    box_bullet(
        doc,
        "Two things are still open and need your judgment as advisors — full detail later in "
        "this document, this is just a heads up.",
    )
    end_box(doc, bg_start, BG_FILL, BG_BORDER)
    doc.add_page_break()

    # ------------------------------------------------------------------
    # TOC
    # ------------------------------------------------------------------
    add_toc(doc)

    doc.add_paragraph(
        "Quick reminder: we're testing whether bond markets actually believed the federal "
        "government when it refused to bail out defaulting states in the 1840s. If defaulted "
        "states kept paying higher interest rates for years afterward, that means the market "
        "believed the promise. It's been about a month since we last talked — this document "
        "starts from scratch, it doesn't assume you remember where we left off."
    )
    doc.add_page_break()

    # ------------------------------------------------------------------
    # Agenda
    # ------------------------------------------------------------------
    doc.add_heading("Agenda", level=1)
    agenda_start = start_box(doc)
    box_heading(doc, "Today's flow")
    for item in [
        "Recap of what we looked into since last meeting",
        "Does a state's default also punish its own city? (Hall's question)",
        "The canal-bond comparison — including one huge finding in Indiana",
        "A better source for our key date",
        "Two things we need your judgment on",
        "A proposed outline for the paper, if we're ready to start writing",
    ]:
        box_bullet(doc, item)
    end_box(doc, agenda_start, AGENDA_FILL, AGENDA_BORDER)
    doc.add_paragraph()

    say_this(
        doc,
        "Thanks for making time — it's been about a month, so before we get to the two things I "
        "actually need from you, I want to walk through everything I looked into. Quick tour: "
        "first a recap of what I dug into, then the city question you raised last time, then a "
        "big finding from the canal-bond comparison, a quick update on our key date, and then the "
        "two decisions I need — and if there's time, a proposed outline for the paper.",
    )

    # ------------------------------------------------------------------
    # What we investigated since last meeting
    # ------------------------------------------------------------------
    doc.add_heading("What We Investigated Since Last Meeting", level=1)
    doc.add_paragraph("Three loose ends from last time — here's what I found on each.")
    say_this(
        doc,
        "So let me start with three things that were still open after our last conversation.",
    )

    doc.add_heading("Bank-held vs. traded bonds investigation", level=2)
    add_bullets(doc, [
        "You flagged that some banks were required to hold state bonds as backing for their own "
        "bank notes — meaning a bond could look \"untraded\" for reasons that have nothing to do "
        "with what investors thought of it. I checked every bond we use for unusually quiet "
        "trading.",
        "One bond (a New York canal bond) is a good match: New York had a real 1838 law requiring "
        "banks to buy and hold bonds like this one, and the timing lines up.",
        "One other quiet bond (an Ohio bond) looks like a different story — trading just moved to "
        "a more popular Ohio bond, not bank-holding.",
        "Two other quiet bonds are still a mystery — not enough evidence either way.",
        "Bottom line: doesn't change any of our results — none of the four are bonds we lean on "
        "for our main findings — but it's an honest caveat to have on record.",
    ])
    say_this(
        doc,
        "First, the bank-holding question you raised — whether some of our quiet bonds were just "
        "sitting in bank vaults instead of trading. Turns out it's a mixed bag: one New York bond "
        "really does look bank-held, one Ohio bond looks more like trading just moved elsewhere, "
        "and two others I honestly can't explain. None of it touches our main results, but I "
        "wanted you to know I checked.",
    )

    doc.add_heading("Alabama railroad bond flag", level=2)
    add_bullets(doc, [
        "A source mentioned Alabama paying bondholders only about 30 cents on the dollar on some "
        "railroad bonds — I wasn't sure if that was part of our 1840s story.",
        "It's not. That's a separate, later episode — Alabama's Reconstruction-era railroad bonds "
        "from the 1870s, thirty years after our data ends. No overlap with our project.",
    ])
    say_this(
        doc,
        "Quick one — I ran into a mention of Alabama paying bondholders 30 cents on the dollar, "
        "and wanted to make sure that wasn't part of our story. It's not — that's a totally "
        "separate railroad-bond episode from the 1870s. Good to rule out, doesn't touch anything.",
    )

    doc.add_heading("Ohio yield-rise claim", level=2)
    add_bullets(doc, [
        "You'd mentioned Ohio's interest rate rose in the 1840s, even though Ohio never "
        "defaulted — I checked whether that meant the \"safe\" state was getting punished too.",
        "It's real, but at the wrong moment. Ohio's rate spiked in early 1842 — that's the "
        "original 1837 crash working through the system, almost a year before the federal "
        "government's actual no-bailout announcement in 1843. By the time the announcement "
        "happened, Ohio's rate had already come back down to normal.",
    ])
    say_this(
        doc,
        "You'd mentioned Ohio's rates rising in the 1840s as maybe a sign the safe state got hit "
        "too. I checked — it's real, but it's the wrong moment. It's the 1837 crash itself, not "
        "the 1843 announcement: Ohio's rate had already spiked and come back down before the "
        "government even made its no-bailout statement. So it's not contagion from the policy, "
        "it's just the original panic.",
    )

    # ------------------------------------------------------------------
    # City contagion
    # ------------------------------------------------------------------
    doc.add_heading("Does a State's Default Punish Its Own City Too?", level=1)
    doc.add_paragraph(
        "This is Hall's question, in plain terms: if a state defaults, does the punishment spread "
        "to something nearby that didn't do anything wrong — specifically, the state's own city? "
        "(One fix: the original notes said \"Pennsylvania vs. another state\" — that was a "
        "mix-up. The real question is Philadelphia, the city, versus Pennsylvania, the state it "
        "sits inside.)"
    )
    say_this(
        doc,
        "Next, the question you raised last time, Hall — does a state's default spread to its own "
        "city, even though the city didn't do anything wrong? I've got four pieces of evidence on "
        "this now.",
    )

    doc.add_heading("Philadelphia vs. Pennsylvania", level=2)
    add_bullets(doc, [
        "Before the announcement, Philadelphia's rate was already a bit lower than Pennsylvania's "
        "— about 1.6 percentage points lower.",
        "After the announcement, that gap widened to about 2.7 percentage points.",
        "Why: Pennsylvania's own rate shot up — toward 20% around its own default — while "
        "Philadelphia's rate barely moved, staying around 5%.",
    ])
    add_chart(
        doc, "chart_city_vs_state.png",
        "Philadelphia's rate (the city, purple) barely reacts to Pennsylvania's default "
        "(the state, blue).",
    )
    say_this(
        doc,
        "First data point: Philadelphia versus Pennsylvania. Before the announcement, the city "
        "was already borrowing a bit cheaper than the state — about a point and a half. After, "
        "that gap widened to about two and a half points, entirely because Pennsylvania's own "
        "rate spiked toward 20% while Philadelphia barely moved.",
    )

    doc.add_heading("Pittsburgh vs. Pennsylvania", level=2)
    add_bullets(doc, [
        "Same test, a second Pennsylvania city: Pittsburgh shows almost the identical gap, about "
        "2.65 percentage points. Same story, second confirmation.",
    ])
    add_chart(
        doc, "chart_pittsburgh_vs_pa.png",
        "Pittsburgh (the city) vs. Pennsylvania (the state) — same pattern as Philadelphia.",
    )
    say_this(
        doc,
        "I wanted to make sure Philadelphia wasn't a fluke, so I checked a second Pennsylvania "
        "city, Pittsburgh — same result, almost exactly the same gap. Two cities, same pattern.",
    )

    doc.add_heading("Cincinnati vs. Ohio — a control", level=2)
    add_bullets(doc, [
        "Then I checked a city in a state that never defaulted — Cincinnati inside Ohio. No gap "
        "at all, basically zero.",
        "That's actually a good sign: the gap only shows up when the state itself is genuinely in "
        "trouble. Ohio was fine, so Cincinnati was fine too.",
    ])
    add_chart(
        doc, "chart_cincinnati_vs_ohio.png",
        "Cincinnati (the city) vs. Ohio (the state) — no gap, because Ohio never defaulted.",
    )
    say_this(
        doc,
        "As a control, I checked Cincinnati inside Ohio — a state that never defaulted. No gap at "
        "all, basically flat. Which is reassuring: it means the city/state gap we're seeing isn't "
        "some generic pattern, it only shows up specifically where the state itself was in real "
        "trouble.",
    )

    doc.add_heading("Second Philadelphia bond", level=2)
    add_bullets(doc, [
        "One gap in the Philadelphia story: our main bond had a missing stretch of data right "
        "after the announcement, so we couldn't see the city's immediate reaction.",
        "Found a second Philadelphia bond that fills exactly that gap — same story, actually a "
        "bigger gap (about 6 percentage points), right in the window we couldn't see before.",
    ])
    say_this(
        doc,
        "One more thing on Philadelphia — our original bond had a hole in the data right after "
        "the announcement, so we couldn't see the city's immediate reaction. I found a second "
        "Philadelphia bond that covers exactly that missing window, and it shows the same "
        "no-punishment story — actually a bigger gap, about 6 points — right when it mattered "
        "most.",
    )

    doc.add_heading("Summary", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "Bottom line: the state gets punished; the city right next to it mostly doesn't. Three "
        "data points from Pennsylvania's own cities agree, and one clean comparison (Cincinnati) "
        "confirms the gap only shows up when the state is actually in trouble."
    )
    run.bold = True
    say_this(
        doc,
        "So overall: state gets punished, city doesn't. Three separate pieces of evidence from "
        "Philadelphia and Pittsburgh all agree, and Cincinnati as a control shows the gap only "
        "appears when the state itself is actually under stress. I think that's a pretty clean "
        "answer to your question.",
    )

    # ------------------------------------------------------------------
    # Canal / robustness
    # ------------------------------------------------------------------
    doc.add_heading("The Canal-Bond Comparison — a Second Check on the Same Question", level=1)
    add_bullets(doc, [
        "Some state debt wasn't a plain loan to the state — it was backed by canal toll money "
        "specifically. We checked those separately, as a second, independent test.",
        "Indiana is the standout: in 1847 Indiana split its canal debt into two slices — one "
        "that got paid first (\"preferred\") and one that got paid only after the first slice was "
        "covered (\"deferred\"). The market priced that difference enormously: the first-in-line "
        "slice traded at 2-4 times the price of the second slice, a roughly 40-percentage-point "
        "interest-rate gap. That's the single biggest number in the whole project.",
        "Checked this wasn't a fluke — it holds up across 146 separate price readings over almost "
        "three years, not a handful of odd days.",
        "Two small data cleanups made along the way, flagged rather than hidden: one bond has a "
        "long gap in its price history (shown as a dotted line, not a straight line, so it "
        "doesn't look like a smooth climb that never happened); one single price reading looks "
        "like a typo in the original 1840s records and is excluded from the chart, though kept in "
        "the raw data.",
    ])
    add_chart(
        doc, "chart_canal_robustness.png",
        "Left: New York's canal bonds. Right: Indiana's first-in-line slice (green) vs. "
        "second-in-line slice (orange) of the same 1847 restructured loan.",
    )
    say_this(
        doc,
        "Now the canal-bond comparison — a second, independent test using a different kind of "
        "state debt, bonds backed specifically by canal toll money. The headline is Indiana: in "
        "1847 they split their canal debt into a slice that got paid first and a slice that got "
        "paid second, and the market priced that difference by about 40 percentage points — by "
        "far the biggest number in the whole project. I stress-tested it the same way I caught a "
        "problem in an earlier finding, and it holds up completely.",
    )

    # ------------------------------------------------------------------
    # Anchor date
    # ------------------------------------------------------------------
    doc.add_heading("A Better Source for Our Key Date", level=1)
    add_bullets(doc, [
        "We've been using February 11, 1843 as the date the federal government made clear it "
        "wouldn't bail states out. Until now that rested only on other historians' books, not an "
        "original document.",
        "I found and read an actual page from the original 1843 Congressional record — a real "
        "debate in Congress on exactly this topic, with the right people talking about the right "
        "issue, sitting right next to another date we'd already confirmed.",
        "One honest caveat: the exact day inside that multi-day debate isn't 100% pinned down, "
        "but everything else lines up.",
    ])
    say_this(
        doc,
        "Quick update on our key date, February 11, 1843. It used to rest only on other "
        "historians' accounts. I tracked down and actually read a page from the original 1843 "
        "Congressional record itself — a real debate on exactly this topic with the right people "
        "in the room. I can't swear to the exact day inside that multi-day debate, but everything "
        "else about it fits.",
    )

    # ------------------------------------------------------------------
    # Two decisions
    # ------------------------------------------------------------------
    doc.add_heading("Two Things We Need Your Judgment On", level=1)
    judgment_start = start_box(doc)
    box_heading(doc, "1. Alabama: default or not?")
    box_bullet(
        doc,
        "We've had Alabama filed under \"defaulted\" this whole time, but growing evidence says "
        "it actually never defaulted — it raised taxes and sold off state bank assets to keep "
        "paying instead. Three separate sources now back this up. We need your call before it "
        "goes into the paper.",
    )
    box_heading(doc, "2. Indiana's second-in-line bondholders — paid less, or paid late?")
    box_bullet(
        doc,
        "Our 40-point Indiana finding is solid on price, but we can't tell from the historical "
        "record whether the second-in-line slice actually stopped getting interest, or just got "
        "paid after the first slice, on a delay. Doesn't change the finding — just how precisely "
        "we describe it.",
    )
    end_box(doc, judgment_start, BOX_FILL, BOX_BORDER)
    p = doc.add_paragraph()
    run = p.add_run(
        "Also worth flagging, not a decision: Cincinnati and the second Philadelphia bond are our "
        "newest findings and haven't been double-checked as hard as everything else yet."
    )
    run.italic = True
    say_this(
        doc,
        "So here's what I actually need from you two. First — do we feel good calling Alabama "
        "\"risky but survived\" instead of \"defaulted\"? I've now got three sources pointing "
        "that way, but I know you weren't sure last time, so I wanted your read before I bake it "
        "into the paper. Second — on the Indiana finding, do you know whether that second-in-line "
        "slice of debt actually stopped paying interest, or just got paid later? It doesn't "
        "change the number, just how precisely I can describe it. And one more thing, not a "
        "decision — Cincinnati and the second Philadelphia bond are our newest findings and "
        "haven't been stress-tested as hard as everything else yet.",
    )

    # ------------------------------------------------------------------
    # Proposed paper structure
    # ------------------------------------------------------------------
    doc.add_heading("If We're Ready to Start Writing — a Proposed Outline", level=1)
    doc.add_paragraph("Not locked in — just a starting point to react to.").italic = True
    outline = [
        "Introduction — the question, tied to the modern EU \"blue bonds\" debate.",
        "Background — the 1837 crash, the 1840s state defaults, the federal refusal to bail "
        "anyone out.",
        "Data and method — old bond prices turned into interest rates, our three state buckets, "
        "and the two different before/after moments we test separately (the 1837 crash itself vs. "
        "the 1843 no-bailout announcement).",
        "Main result — does the no-bailout promise show up in state interest rates: Pennsylvania, "
        "Ohio, Alabama, Indiana, New York.",
        "City result — does a state's default punish its own city too.",
        "Robustness — the canal-bond comparison and the Indiana finding.",
        "Discussion — the market drawing careful lines instead of panicking blindly, tied back to "
        "the EU parallel.",
        "Conclusion and limitations — the gaps and open questions we're upfront about.",
    ]
    for item in outline:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
    say_this(
        doc,
        "If you think we're in good shape, here's roughly how I'm picturing the paper — let me "
        "know if this order makes sense or if you'd restructure it. Nothing here is locked in, "
        "this is just a starting point to react to.",
    )

    # ------------------------------------------------------------------
    # Closing
    # ------------------------------------------------------------------
    doc.add_heading("Wrapping Up", level=1)
    say_this(
        doc,
        "So to sum up where we landed: I need a yes-or-no on Alabama, and your best guess on the "
        "Indiana coupon question, and if the paper outline looks right, I'll start drafting. I'll "
        "write up any changes from today and send it your way — thanks again for the time, this "
        "was really helpful.",
    )

    out_path = OUTPUT_DIR / "meeting_prep_v2.docx"
    doc.save(str(out_path))
    print(f"saved -> {out_path}")
    print("\nEmbedded charts:")
    for c in embedded_charts:
        print(f"  - {c}")
    if missing_charts:
        print("\nMISSING charts (placeholder inserted):")
        for c in missing_charts:
            print(f"  - {c}")


if __name__ == "__main__":
    main()
