"""Build fact_check.docx / fact_check.pdf -- a human-readable, bulleted
fact-check sheet for the 1840s State Debt Credibility project.

Same facts as output/FACTS_LAB_REPORT.md, but rewritten to actually read
(short bullets, plain sentences, no dense tables), in a warm serif
("Claude"-style) typeface, with the project's charts pulled in so the
whole thing is one self-contained fact-check document.

Every number is copied from FACTS_LAB_REPORT.md, which was recomputed
against the source CSVs on 2026-09-02.

Rendered to PDF via Pages (osascript) after this writes the .docx.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "output"
CHART_DIR = OUTPUT_DIR  # use the *_clean.png versions (no dense footnotes)

BODY = "Georgia"
INK = RGBColor(0x33, 0x30, 0x2B)          # warm charcoal
MUTED = RGBColor(0x86, 0x7F, 0x74)        # warm gray, for source/trace notes
ACCENT_RGB = RGBColor(0xB2, 0x50, 0x33)   # Claude-ish warm rust
ACCENT_HEX = "B25033"
RULE_HEX = "E3D4C8"                       # pale warm rule under H1

embedded = []
missing = []


def _rule_below(style, hex_color, sz="6"):
    pPr = style.element.get_or_add_pPr()
    b = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    b.append(bottom)
    pPr.append(b)


def _set_font(run, name=BODY, size=None, color=None, italic=False, bold=False):
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    run.italic = italic
    run.bold = bold


def para(doc, text, size=10.5, color=INK, italic=False, bold=False, after=6, before=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_font(r, size=size, color=color, italic=italic, bold=bold)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    return p


def bullets(doc, items, size=10.5):
    """items: str, or (lead, rest) -> lead rendered bold."""
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            lead, rest = it
            r1 = p.add_run(lead)
            _set_font(r1, size=size, color=INK, bold=True)
            r2 = p.add_run(rest)
            _set_font(r2, size=size, color=INK)
        else:
            r = p.add_run(it)
            _set_font(r, size=size, color=INK)
        p.paragraph_format.space_after = Pt(3)


def source_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_font(r, size=8.5, color=MUTED, italic=True)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(1)


def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        _set_font(r, size=14, color=ACCENT_RGB, bold=True)
    h.paragraph_format.keep_with_next = True
    return h


def h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        _set_font(r, size=11.5, color=ACCENT_RGB, bold=True)
    h.paragraph_format.keep_with_next = True
    return h


def chart(doc, filename, caption, width_in=6.1):
    path = CHART_DIR / filename
    if path.exists():
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before = Pt(6)
        pp.paragraph_format.space_after = Pt(2)
        pp.add_run().add_picture(str(path), width=Inches(width_in))
        cp = doc.add_paragraph()
        cr = cp.add_run(caption)
        _set_font(cr, size=8.5, color=MUTED, italic=True)
        cp.paragraph_format.space_after = Pt(12)
        embedded.append(filename)
    else:
        para(doc, f"[chart missing: {filename}]", color=RGBColor(0xB0, 0, 0), bold=True)
        missing.append(filename)


def page_numbers(doc):
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    a = OxmlElement("w:fldChar"); a.set(qn("w:fldCharType"), "begin")
    b = OxmlElement("w:instrText"); b.set(qn("xml:space"), "preserve"); b.text = "PAGE"
    c = OxmlElement("w:fldChar"); c.set(qn("w:fldCharType"), "end")
    run._r.append(a); run._r.append(b); run._r.append(c)
    _set_font(run, size=8.5, color=MUTED)


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    hs1 = doc.styles["Heading 1"]
    hs1.font.name = BODY
    hs1.paragraph_format.space_before = Pt(16)
    hs1.paragraph_format.space_after = Pt(6)
    _rule_below(hs1, RULE_HEX, sz="6")

    hs2 = doc.styles["Heading 2"]
    hs2.font.name = BODY
    hs2.paragraph_format.space_before = Pt(11)
    hs2.paragraph_format.space_after = Pt(3)

    # ---- Title ----
    t = doc.add_heading("Fact-Check Sheet — 1840s State Debt Credibility Project", level=0)
    for r in t.runs:
        _set_font(r, size=19, color=ACCENT_RGB, bold=True)
    t.paragraph_format.space_after = Pt(4)
    para(
        doc,
        "Every figure here was recomputed from the source data on 2 September 2026. Where a "
        "number is rounded or a prose version elsewhere disagrees slightly, that's flagged. "
        "This is the sheet to hold the paper against — dry by design, but written to read.",
        size=10, color=MUTED, italic=True, after=10,
    )

    # ---- The question ----
    h1(doc, "The question we're testing")
    bullets(doc, [
        "Did the transatlantic bond market believe the U.S. federal government's early-1840s "
        "refusal to bail out defaulting states — treat it as a permanent commitment rather "
        "than a bluff?",
        "The test: if the market believed it, states that defaulted should have kept paying "
        "noticeably higher interest for years afterward, states that stayed safe shouldn't "
        "have, and states that got scared but paid up should land somewhere in between.",
        "We build the interest rates ourselves from historical bond prices — the records give "
        "prices, not yields — covering the 1830s through the 1850s.",
        "Why it matters now: the EU's current \"blue bonds\" debate is the same dilemma — does "
        "a credible no-bailout rule actually force fiscal discipline? (Hanno Lustig, March "
        "2026.)",
    ])

    # ---- The data ----
    h1(doc, "The data we're working from")
    bullets(doc, [
        ("Bond prices — ", "the EH.net Early U.S. Securities Prices database: the New York "
         "exchange file and the Philadelphia exchange file. Parsed into "
         "new_york_state_debt_prices.csv (3,492 weekly rows, 194 securities) and "
         "philadelphia_state_debt_prices.csv (3,147 rows, 37 securities)."),
        ("The codebook — ", "Securities Index.xls maps each bond code to its name, coupon, and "
         "maturity year. It has no issue-date field for anything, which is why we can't date "
         "the Indiana canal bonds."),
        ("Not used — ", "the Bayley federal-loans document is in the repo but turned out to be "
         "a history of federal Treasury borrowing, not state-bond detail. Other cities' "
         "exchange files (Baltimore, Boston, etc.) were never downloaded."),
    ])
    h2(doc, "The anchor date for the \"no-bailout signal\"")
    bullets(doc, [
        "We date it to February 1843 — the 11 February 1843 House floor debate on assuming "
        "state debts (27th Congress, 3rd session, under Tyler), reinforced by the U.S. consul "
        "at The Hague telling European bankers in April–July 1843 that the federal government "
        "would not be held responsible \"for any default, actual or eventual.\"",
        "Caveat to carry into the paper: two secondary sources plus a genuine primary-source "
        "read of the right Congressional Globe volume (right people, right topic, right week) "
        "— but the exact day, Feb 11 vs. an adjacent day of the same multi-day debate, isn't "
        "pinned down.",
        "The earlier 1839–40 assumption episode is deliberately not the anchor — it predates "
        "the actual defaults.",
    ])

    # ---- Method ----
    h1(doc, "How prices become interest rates")
    bullets(doc, [
        "For Pennsylvania, Ohio, and New York we use yield-to-maturity — there's a real "
        "maturity year to work from.",
        "For Alabama and Indiana we use current yield (coupon ÷ price) — their codebook "
        "entries have no usable maturity date. This is a genuine inconsistency between states; "
        "we keep it visible in the data rather than hide it.",
        ("The active-default override — ", "yield-to-maturity assumes a bond is repaid in "
         "full, on time. That's plainly wrong for a state currently in default, so for any "
         "date inside a state's known default window we switch that bond to current yield. "
         "This is what fixed Pennsylvania's numbers (below). Pennsylvania's window: Aug 1842 "
         "to Feb 1845. Indiana's: Jan 1841 onward. Alabama is deliberately not on the list — "
         "it raised taxes and liquidated its state bank to keep paying, and never defaulted."),
        "Bonds trading within a year of their own maturity, or past it, are flagged and "
        "dropped from the trend lines — the math gets unstable near maturity. The rows stay "
        "in the file, just flagged.",
    ])

    # ---- Buckets ----
    h1(doc, "The three buckets")
    bullets(doc, [
        ("Pennsylvania — defaulted (confirmed). ", "Bonds S-2240 / S-2250 / S-2270 / S-2330 / "
         "S-2410. Suspended interest August 1842, handed creditors paper instead of cash, "
         "resumed 1845. The marquee case."),
        ("Indiana — defaulted / restructured (confirmed). ", "Bonds S-0510 / S-0540 (the "
         "original general loan). Defaulted January 1841; the 1846–47 Butler Bill handed the "
         "Wabash & Erie Canal to creditors for half the debt."),
        ("Ohio — safe (confirmed). ", "Bonds S-2100 / S-2110 / S-2080 / S-2010. Never "
         "defaulted; took a hard hit in the 1842 panic and fully recovered by 1844."),
        ("New York — risky but survived (bucket confirmed; the premium is weak in the data). ",
         "Bond S-1650. Real railroad and bank-failure stress, never defaulted."),
        ("Alabama — risky but survived (still needs Hall / Sargent's sign-off). ", "Bonds "
         "S-0030 / S-0040. Three independent secondary sources say it didn't default; the "
         "advisors haven't personally confirmed the reclassification."),
    ])

    # ---- Results ----
    h1(doc, "What we found — the numbers")

    h2(doc, "Main comparison, after the no-bailout signal (April 1843 on)")
    bullets(doc, [
        "Pennsylvania (defaulted): 8.76% average — over a short window that runs out in "
        "January 1845 (88 dates).",
        "Indiana (defaulted): 13.87% average, through December 1848 (203 dates).",
        "Alabama (risky / survived): 7.01%, through mid-1853 (142 dates).",
        "Ohio (safe): 5.78%, through the end of 1853 (444 dates).",
        "New York (risky / survived): 5.37%, through January 1848 (117 dates).",
        ("Reading: ", "the two clear defaulters carry a premium — Indiana a big one; Alabama "
         "sits a bit above safe Ohio; New York sits at or below it. The pattern holds for "
         "Pennsylvania and Indiana and is weak for New York."),
    ])
    source_note(doc, "Source: primary_yields.csv, via scripts/calculate_yields.py — usable rows, averaged by date.")
    chart(doc, "chart_policy_short_medium_clean.png",
          "Policy window, each state truncated to its own real coverage. Pennsylvania sits "
          "high but its line is short; Ohio and New York run low; Alabama in between.")
    chart(doc, "chart_policy_long_term_clean.png",
          "The long run — only Ohio and Alabama have dense decade-long data. Alabama holds a "
          "~1–2 point gap over safe Ohio that narrows but never closes.")

    h2(doc, "The Pennsylvania correction — important")
    bullets(doc, [
        "Before the active-default override, Pennsylvania's post-signal average came out at "
        "19.81% — a 13-point gap over Ohio. That was an artifact: yield-to-maturity was "
        "assuming a defaulted bond would still be repaid at par.",
        "After the override it's 8.77% — a ~2-point gap over Ohio, not 13.",
        ("Bottom line: ", "the persistence finding survives, but it's a modest 2 points, not "
         "a dramatic 13. The old 13-point number must not appear anywhere in the paper."),
    ])
    source_note(doc, "Source: primary_yields.csv vs primary_yields_before_default_override.csv. "
                     "221 rows had their yield changed by the override.")

    h2(doc, "Alabama vs. Ohio, year by year")
    bullets(doc, [
        "1843: 8.23% vs 8.18% — basically identical (+0.04 points).",
        "1845: 7.11% vs 6.31% (+0.81).   1847: 8.26% vs 6.22% (+2.04).   1848: 8.12% vs 6.11% (+2.01).",
        "1850: 6.21% vs 4.83% (+1.38).   1852–53: ~5.5% vs ~4.4% (~+1.0).",
        ("Reading: ", "Alabama opens level with safe Ohio right after the signal, opens a "
         "~2-point gap by 1847–48, then narrows but never fully closes. A clean \"landed in "
         "between\" pattern — arguably cleaner than New York's."),
        ("Rounding note: ", "PROJECT_CONTEXT prose says the 1843 gap is +0.05 and 1847 Alabama "
         "is 8.22%; recomputed they're +0.04 and 8.26%. Rounding only — story unchanged."),
    ])
    source_note(doc, "Source: primary_yields.csv, full-calendar-year averages.")

    h2(doc, "New York vs. Ohio, and Indiana's own bonds")
    bullets(doc, [
        "New York runs 5.0–5.6% and sits below safe Ohio for essentially the whole window, "
        "ticking up to 7.0% only in a single January 1848 reading. On this bond, New York "
        "reads more like \"safe\" than \"risky but survived.\"",
        "Indiana's general-obligation bonds run 12.6–16.6% year by year, 1843–1848 — a large, "
        "sustained premium (roughly 6–10 points over Ohio). The strongest defaulted-bucket "
        "signal we have.",
    ])

    h2(doc, "The 1837 panic, for context")
    bullets(doc, [
        "Pennsylvania's bonds fell from the 90–110 range to 37–40 by August–September 1842, "
        "right at the default.",
        "Ohio's bonds bottomed in a tight four-week window, March–April 1842, with yields "
        "spiking to 11.5–16%. Then they recovered fully by 1844 — before the 1843 policy "
        "signal. So Ohio's yield spike is the panic, not the policy.",
    ])
    source_note(doc, "Source: raw price files; output/ohio_yield_check.csv.")
    chart(doc, "chart_panic_window_clean.png",
          "The panic window — every state collapses in 1841–42. Pennsylvania and the "
          "defaulters don't recover; Ohio and New York snap back by 1843–44.")

    h2(doc, "Does a state's default spill onto its own city?")
    bullets(doc, [
        ("Philadelphia vs. Pennsylvania: ", "before the signal the city was already ~1.6 "
         "points cheaper than the state; after, in the overlap window where we have both, "
         "2.68 points. The whole widening is Pennsylvania's own rate climbing — Philadelphia's "
         "barely moved."),
        ("Pittsburgh vs. Pennsylvania: ", "a 2.65-point gap post-signal — essentially the "
         "same as Philadelphia. A second Pennsylvania city, same result."),
        ("Cincinnati vs. Ohio: ", "basically no gap (−0.11 points). But Ohio never defaulted, "
         "so this isn't a third confirmation — it's the baseline: when the state's fine, the "
         "city tracks it."),
        ("New York City vs. New York State: ", "also near zero (≈ −0.19). Same story as "
         "Cincinnati — a safe state, no city gap."),
        ("Reading: ", "the city/state gap only opens when the state itself is in trouble. The "
         "market told the two credits apart."),
    ])
    source_note(doc, "Source: scripts/compare_city_vs_state*.py; city_vs_state*.csv; nyc_brooklyn_check.csv.")
    chart(doc, "chart_city_vs_state_clean.png",
          "Philadelphia (city) vs. Pennsylvania (state). The state's line spikes toward 20%; "
          "the city's stays flat around 5%.")
    chart(doc, "chart_pittsburgh_vs_pa_clean.png",
          "Pittsburgh (city) vs. Pennsylvania (state) — the same ~2.65-point gap as "
          "Philadelphia.")
    chart(doc, "chart_cincinnati_vs_ohio_clean.png",
          "Cincinnati (city) vs. Ohio (state) — no gap, because Ohio was never in trouble. "
          "The baseline contrast.")

    h2(doc, "The second Philadelphia bond — closes the biggest gap in that finding")
    bullets(doc, [
        "The main Philadelphia bond has a 343-day hole in its data covering exactly the "
        "immediate post-signal months. A second Philadelphia bond (C-1260) fills that hole.",
        "In that window: Philadelphia ~5.6% vs. Pennsylvania 11.5–11.9% — a 5.95 to "
        "6.28-point gap, right when it mattered.",
        ("So the no-spillover finding is if anything stronger than first reported. ", "The "
         "earlier \"5.4 points\" figure was a miscount and should be replaced with these."),
    ])
    source_note(doc, "Source: output/philadelphia_second_bond_check.csv. Not yet folded into the protected comparison script.")

    h2(doc, "The canal bonds and the Indiana tranches — the robustness check")
    bullets(doc, [
        "A separate test using bonds backed specifically by canal toll revenue.",
        ("Indiana's 1847 restructuring split its canal debt into \"preferred\" (paid first) "
         "and \"deferred\" (paid after) slices. ", "The market priced that split by ~40 "
         "percentage points — deferred averaged 58% current yield, preferred ~18%. The "
         "largest gap anywhere in the project."),
        "Stress-tested the same way we caught the Pennsylvania problem: it holds across 146 "
        "observations over ~3 years, it's not a handful of odd days, and no maturity-math "
        "artifact is possible here.",
        "All this data is 1850–1853 though — years after the policy signal — so it speaks to "
        "seniority pricing, not the 1843 question.",
        ("Two data-quality catches, flagged not hidden: ", "one New York canal bond has a "
         "4.5-year gap (drawn as a broken line, not a smooth climb); one single New York "
         "price of 160.00 looks like a transcription error in the original records (kept in "
         "the file, left off the chart)."),
    ])
    source_note(doc, "Source: scripts/compare_canal_robustness.py; canal_robustness_yields.csv; indiana_tranche_sanity_check.csv.")
    chart(doc, "chart_canal_robustness_clean.png",
          "Left: New York's canal bonds. Right: Indiana's preferred (green) vs. deferred "
          "(orange) tranches of the 1847 restructured loan — a ~40-point gap.")

    h2(doc, "Thinly-traded bonds")
    bullets(doc, [
        "Four bonds trade unusually thin: Alabama's S-0040, Ohio's S-2010, New York's canal "
        "bond S-1750, and Ohio's canal bond S-2190 (only 2 observations, both 1825 — "
        "unusable).",
        "New York's 1838 Free Banking Act (banks had to hold state bonds) is the best fit for "
        "S-1750's sparsity. The others we can't fully explain.",
        ("None of these carry a headline number, ", "so this doesn't change any result."),
    ])
    source_note(doc, "Source: output/trade_density.csv.")

    # ---- Caveats ----
    h1(doc, "Every caveat on record")
    para(doc, "Severity: H = affects a headline number.  M = affects how far a claim "
              "generalizes.  L = affects wording only.", size=9.5, color=MUTED, italic=True, after=4)
    bullets(doc, [
        ("[H] ", "Pennsylvania / Ohio / New York use yield-to-maturity; Alabama and Indiana "
         "use current yield. The two diverge most when prices are far from par — exactly "
         "mid-panic. This bit the Alabama–Ohio 1843 row once; now fixed."),
        ("[H] ", "Pennsylvania's post-signal data runs out around January 1845. Its "
         "\"persistence\" claim is really a ~20-month test, not a decade."),
        ("[H] ", "New York's usable bond (S-1650) doesn't start until July 1842 — it missed "
         "the worst of New York's 1839–42 stress. Its \"no premium over Ohio\" result is "
         "substantially a data-coverage artifact, not clean evidence New York was safe."),
        ("[H] ", "The Alabama reclassification (defaulted → risky-but-survived) is still not "
         "personally confirmed by Hall or Sargent."),
        ("[M] ", "New York's bond data ends December 1848 — no long-run picture for New York."),
        ("[M] ", "Alabama and Indiana have no codebook maturity, so we can't even run the "
         "near-maturity check for them — a real blind spot."),
        ("[M] ", "The Philadelphia main-bond gap means the immediate post-signal reaction was "
         "only recoverable via the second bond, which isn't yet folded into the protected "
         "comparison script."),
        ("[M] ", "The city comparison is a handful of pairs — two Pennsylvania cities for the "
         "stress case, two safe-state cities as baseline. Not \"cities in general.\""),
        ("[M] ", "The default-period dates that drive the override are secondary-source, not "
         "primary."),
        ("[M] ", "The \"why the city stayed solvent\" explanation is inference from legal / "
         "financial structure plus our own price data — no single contemporary source states "
         "it outright, and only Philadelphia was researched."),
        ("[L] ", "The 11 February 1843 anchor date isn't pinned to the exact day."),
        ("[L] ", "The Indiana canal-tranche data is all 1850–53 — it tests seniority pricing, "
         "not the 1843 signal."),
        ("[L] ", "Whether Indiana's \"deferred\" tranche stopped paying interest or just paid "
         "late is unresolved — affects the wording of the 40-point figure, not the figure."),
        ("[L] ", "Indiana canal-bond issuance timing can't be determined — no issue-date "
         "field, no legislative rollout schedule."),
        ("[L] ", "One New York canal price (160.00) and one Indiana first print (50.00) are "
         "likely transcription errors — kept in the files, off the charts."),
        ("[L] ", "primary_yields.csv has 3,711 rows. An old note said 3,441 — that was before "
         "two Pennsylvania bonds (S-2330, S-2410) were added; 3,441 + 270 = 3,711."),
    ])

    # ---- Open items ----
    h1(doc, "What's still open")
    h2(doc, "Needs the advisors")
    bullets(doc, [
        "Alabama: defaulted, or risky-but-survived? Three sources say the latter; need a "
        "yes / no.",
        "Indiana's deferred tranche: interest suspended, or just subordinated-but-paid? "
        "Changes wording, not the finding.",
        "Indiana canal payment rule: strict waterfall, or proportional split? No source "
        "states it; the prices suggest both classes took a hit.",
    ])
    h2(doc, "Needs more research (already tried)")
    bullets(doc, [
        "Indiana canal-bond issue dates — the codebook, the market data, and the standard "
        "histories all come up empty.",
        "A source that directly states why Philadelphia's credit held while Pennsylvania's "
        "collapsed — we have the structural facts, not a single quotable line.",
        "Folding the second Philadelphia bond into the protected comparison script (~15 "
        "minutes, not advisor work).",
        "Reading the exact day off the Congressional Globe (needs a human to clear a CAPTCHA).",
    ])
    h2(doc, "Can't be resolved")
    bullets(doc, [
        "Whether New York's general-obligation bond dipped in 1839–42 — it simply wasn't "
        "quoted then.",
        "The two thinniest-traded bonds (Alabama S-0040, Ohio canal S-2190) — too few "
        "records survive.",
    ])

    # ---- Traceability ----
    h1(doc, "Where every number comes from")
    bullets(doc, [
        ("scripts/calculate_yields.py → ", "primary_yields.csv (3,711 rows) + "
         "primary_yields_before_default_override.csv. The core pipeline — all state-level "
         "yields, the truncation flags, the active-default override."),
        ("scripts/compare_city_vs_state.py → ", "city_vs_state_yields.csv + "
         "chart_city_vs_state.png. Philadelphia vs. Pennsylvania."),
        ("scripts/compare_city_vs_state_cincinnati_pittsburgh.py → ", "city_vs_state_cincinnati.csv, "
         "city_vs_state_pittsburgh.csv + their charts."),
        ("scripts/compare_canal_robustness.py → ", "canal_robustness_yields.csv + "
         "chart_canal_robustness.png. New York canal bonds and the Indiana tranches."),
        ("scripts/build_yield_charts.py → ", "chart_panic_window / chart_policy_short_medium / "
         "chart_policy_long_term. The three-tier chart set."),
        ("Diagnostic files (one-off, no saved script): ", "trade_density.csv, "
         "ohio_yield_check.csv, nyc_brooklyn_check.csv, philadelphia_second_bond_check.csv, "
         "philadelphia_county_check.csv, indiana_tranche_sanity_check.csv."),
        ("The full dry version, with the exact tables: ", "output/FACTS_LAB_REPORT.md."),
    ])
    para(doc, "")
    para(doc,
         "All five pipeline scripts were re-run on 2 September 2026 and reproduce their "
         "outputs byte-for-byte. Every figure above traces to one of the files listed here; "
         "the only claims that don't are the historical dates and institutional facts (the "
         "anchor date, the default periods, the bucket evidence, the Butler Bill chronology, "
         "the municipal-law points), which rest on secondary sources and are flagged where "
         "they appear.",
         size=9.5, color=MUTED, italic=True)

    page_numbers(doc)

    out = OUTPUT_DIR / "fact_check.docx"
    doc.save(str(out))
    print(f"saved -> {out}")
    print("charts embedded:")
    for c in embedded:
        print("  -", c)
    if missing:
        print("MISSING:")
        for c in missing:
            print("  -", c)


if __name__ == "__main__":
    main()
