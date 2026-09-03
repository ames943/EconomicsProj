"""Build meeting_prep_v4.docx -- personal-use-only, builds on
meeting_prep_final_private.docx (v3-private).

Not shown to anyone (per the earlier task: "This doc is for ME ONLY").
Charts are presented separately from presentation_charts/; this docx
keeps its embedded charts only as Amey's own reference while prepping.

Changes from meeting_prep_final_private.docx (2026-09-02):
1. New section "Why Didn't the City Itself Default?" -- inserted right
   after the existing city-contagion "Summary" subsection. Explains the
   legal/financial reason the city stayed solvent (separate corporation,
   separate revenue base), kept visibly distinct from the market-pricing
   "no-spillover" finding above it. Written explanation subsection +
   SAY THIS block, same visual treatment as the rest of the doc.
2. New section "The Indiana Payment Order -- Answering Hall's Question"
   -- inserted right after the existing "Canal-Bond Comparison" section
   (which it directly follows up on). Priority mapping of the five
   Indiana codes, the sub-ranking inside "preferred", what the payment
   rule / issuance timing genuinely can't be determined, and what the
   price data itself implies. Written explanation + SAY THIS block.
3. Cheat-sheet page: a third, smaller line added below the two decision
   blocks (the two decisions themselves are unchanged) pointing to the
   two new sections and their page numbers.

Everything else (page-number footer, note-to-self styling, bordered-
paragraph boxes, accent colors) is carried over from
build_meeting_prep_final_private.py unchanged.

Boxes are still built from bordered paragraphs, not tables (see prior
passes' notes on why -- a table-based box silently clipped its last line
under Pages' renderer).
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "output"
CHART_DIR = OUTPUT_DIR / "charts_v3_cropped"

INK = RGBColor(0x0B, 0x0B, 0x0B)
MUTED = RGBColor(0x52, 0x51, 0x4E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

ACCENT = "1B4B66"
ACCENT_RGB = RGBColor(0x1B, 0x4B, 0x66)
ACCENT_TINT = "E9F1F5"

DECISION_BORDER = "C97A26"
DECISION_FILL = "FFF3E0"
DECISION_TEXT = RGBColor(0x8A, 0x51, 0x14)

BG_STRIPE = "9C8B6E"
BG_FILL = "F2EFE9"

# Cheat-sheet cross-reference page numbers for the two new sections.
# Set after building the doc once and checking where each section's H1
# lands in the rendered PDF (see build note / report for 2026-09-02).
CITY_DEFAULT_PAGE = 8      # "Why Didn't the City Itself Default?"
INDIANA_ORDER_PAGE = 9     # "The Indiana Payment Order — Answering Hall's Question"

embedded_charts = []
missing_charts = []


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(4)


def muted_note_bullet(doc, text):
    """A 'note to self' bullet -- italic + muted gray, no border. Distinct
    from a normal factual bullet and from a SAY THIS box."""
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(4)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_chart(doc, filename, caption):
    path = CHART_DIR / filename
    if path.exists():
        pic_p = doc.add_paragraph()
        pic_p.paragraph_format.space_before = Pt(4)
        pic_p.paragraph_format.space_after = Pt(0)
        run = pic_p.add_run()
        run.add_picture(str(path), width=Inches(6.0))
        embedded_charts.append(str(path))
        add_caption(doc, caption)
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[chart not found — check output/: {filename}]")
        run.bold = True
        run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
        missing_charts.append(filename)


def set_style_bottom_rule(style, color_hex, size="6"):
    pPr = style.element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_numbers(doc):
    """Bottom-center page numbers on every page except the very first
    (the cheat-sheet page)."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    # section.first_page_footer is left untouched (its default single empty
    # paragraph), so page 1 shows no number.
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_cheat_sheet(doc):
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(70)
    title_p.paragraph_format.space_after = Pt(50)
    run = title_p.add_run("IF NOTHING ELSE —\nGET THESE TWO ANSWERS")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = ACCENT_RGB

    def block(question, sub):
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(30)
        qp.paragraph_format.space_after = Pt(6)
        qp.paragraph_format.left_indent = Pt(24)
        qp.paragraph_format.right_indent = Pt(24)
        r = qp.add_run(question)
        r.bold = True
        r.font.size = Pt(19)
        r.font.color.rgb = DECISION_TEXT

        sp = doc.add_paragraph()
        sp.paragraph_format.left_indent = Pt(24)
        sp.paragraph_format.right_indent = Pt(24)
        r2 = sp.add_run(sub)
        r2.italic = True
        r2.font.size = Pt(13)
        r2.font.color.rgb = MUTED

    block(
        "Decision 1 — Alabama: defaulted, or risky-but-survived?",
        "3 sources now support reclassifying it; not yet confirmed by either of them.",
    )
    block(
        "Decision 2 — Indiana's deferred tranche: did it stop paying interest, or just get paid later?",
        "Doesn't change the 40pp finding, only its wording.",
    )

    extra = doc.add_paragraph()
    extra.paragraph_format.space_before = Pt(44)
    extra.paragraph_format.left_indent = Pt(24)
    extra.paragraph_format.right_indent = Pt(24)
    r = extra.add_run(
        "Also ready to share if it comes up: why Philadelphia itself never defaulted "
        f"(p. {CITY_DEFAULT_PAGE}), and the Indiana payment-order mapping "
        f"(p. {INDIANA_ORDER_PAGE}) — both now answered."
    )
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED

    doc.add_page_break()


# --- Bordered-paragraph box ---

def start_box(doc):
    return len(doc.paragraphs)


def end_box(doc, start_index, fill_hex, border_hex, indent_pt=14, label_fill=None,
            label_text_white=False, stripe=False):
    paras = doc.paragraphs[start_index:]
    n = len(paras)
    for i, p in enumerate(paras):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        if stripe:
            edge_specs = [("left", "48")]
        else:
            edge_specs = [("left", "16"), ("right", "16")]
            if i == 0:
                edge_specs.append(("top", "16"))
            if i == n - 1:
                edge_specs.append(("bottom", "16"))
        for edge, sz in edge_specs:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), sz)
            el.set(qn("w:space"), "8")
            el.set(qn("w:color"), border_hex)
            pBdr.append(el)
        pPr.append(pBdr)
        this_fill = label_fill if (label_fill and i == 0) else fill_hex
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), this_fill)
        pPr.append(shd)
        p.paragraph_format.left_indent = Pt(indent_pt + (10 if stripe else 0))
        p.paragraph_format.right_indent = Pt(indent_pt)
        p.paragraph_format.space_before = Pt(6 if i > 0 else 7)
        p.paragraph_format.space_after = Pt(6 if i < n - 1 else 7)
        if label_fill and i == 0 and label_text_white:
            for run in p.runs:
                run.font.color.rgb = WHITE


def box_heading(doc, text, size=12, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def box_bullet(doc, text):
    p = doc.add_paragraph()
    p.add_run("•  " + text)
    return p


def box_paragraph(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    return p


def say_this(doc, text):
    start = start_box(doc)
    label = doc.add_paragraph()
    run = label.add_run("SAY THIS")
    run.bold = True
    run.font.small_caps = True
    run.font.size = Pt(10)
    body = doc.add_paragraph()
    run2 = body.add_run(text)
    run2.italic = True
    end_box(doc, start, ACCENT_TINT, ACCENT, label_fill=ACCENT, label_text_white=True)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(10)


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)

    h1 = doc.styles["Heading 1"]
    h1.font.color.rgb = ACCENT_RGB
    h1.font.size = Pt(17)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    set_style_bottom_rule(h1, ACCENT, size="6")

    h2 = doc.styles["Heading 2"]
    h2.font.color.rgb = ACCENT_RGB
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)

    # ------------------------------------------------------------------
    # Page 1: cheat sheet
    # ------------------------------------------------------------------
    add_cheat_sheet(doc)

    # ------------------------------------------------------------------
    # Title
    # ------------------------------------------------------------------
    title = doc.add_heading("1840s State Debt Credibility — Progress Since Last Meeting", level=0)
    title.paragraph_format.space_after = Pt(10)

    # ------------------------------------------------------------------
    # Background (bulleted, distinct stripe panel -- read silently)
    # ------------------------------------------------------------------
    doc.add_heading("Background — What Is This Project, Actually", level=1)
    bg_start = start_box(doc)
    box_paragraph(
        doc,
        "(Read this before the meeting — it's for you, not something to say out loud.)",
        italic=True, bold=True,
    )

    box_heading(doc, "The basic story", size=11)
    for item in [
        "States borrowed heavily from British and Dutch investors in the 1830s.",
        "A financial crash in 1837 wrecked state tax revenue.",
        "By the early 1840s, nine states plus the Florida Territory defaulted.",
        "The federal government refused to bail any of them out.",
        "Our evidence: old bond price records from 1830s-1850s newspapers and exchanges, "
        "turned into interest rates.",
    ]:
        box_bullet(doc, item)

    box_heading(doc, "The question we're testing", size=11)
    for item in [
        "Did investors actually believe the federal government meant it?",
        "If they believed it, defaulted states should have kept paying higher interest rates "
        "for years.",
        "If they didn't believe it, rates should have bounced back to normal quickly.",
    ]:
        box_bullet(doc, item)

    box_heading(doc, "Why it matters today", size=11)
    for item in [
        "Economist Hanno Lustig's March 2026 post \"Blue Bonds for Europe\" says the EU faces "
        "this same choice now.",
        "Same question, 200 years apart: does a credible \"no bailout\" promise actually make "
        "governments more careful with money?",
    ]:
        box_bullet(doc, item)

    box_heading(doc, "The three groups of states", size=11)
    for item in [
        "Defaulted: Pennsylvania and Indiana for sure, likely Alabama too (still to be confirmed).",
        "Scared but paid everyone back: New York.",
        "Never really at risk: Ohio — our safe baseline.",
    ]:
        box_bullet(doc, item)

    box_heading(doc, "The city side-question", size=11)
    for item in [
        "Hall asked: does a state's default also punish the city inside it, like Philadelphia "
        "inside Pennsylvania?",
        "Even though the city itself did nothing wrong and kept paying its own debts.",
    ]:
        box_bullet(doc, item)

    box_heading(doc, "What we found", size=11)
    box_bullet(
        doc,
        "States that defaulted kept paying more in interest for years afterward — the market "
        "didn't forgive and forget.",
    )
    box_bullet(
        doc,
        "The cities inside those states did NOT get punished — investors could tell the "
        "difference between a state's credit and its own city's credit, even sharing a name and "
        "a place.",
    )
    box_bullet(
        doc,
        "One huge extra finding worth flagging: in Indiana, one slice of a restructured loan got "
        "paid before another slice (a \"tranche\" just means one slice of a loan with its own "
        "place in line) — and the market priced that difference by more than anything else we "
        "found in this whole project.",
    )
    box_bullet(
        doc,
        "Two things are still open and need your judgment as advisors — full detail later in "
        "this document, this is just a heads up.",
    )
    end_box(doc, bg_start, BG_FILL, BG_STRIPE, stripe=True)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(14)

    # ------------------------------------------------------------------
    # Recap
    # ------------------------------------------------------------------
    doc.add_paragraph(
        "Quick reminder: we're testing whether bond markets actually believed the federal "
        "government when it refused to bail out defaulting states in the 1840s. If defaulted "
        "states kept paying higher interest rates for years afterward, that means the market "
        "believed the promise."
    )

    say_this(
        doc,
        "It's been about a month since we last talked, so I want to walk through everything from "
        "scratch — nothing here assumes you remember where we left off. By the end there are two "
        "decisions I'll need from you. Let's dive in.",
    )

    # ------------------------------------------------------------------
    # Agenda
    # ------------------------------------------------------------------
    doc.add_heading("Agenda", level=1)
    agenda_start = start_box(doc)
    box_heading(doc, "TODAY'S FLOW", size=10, color=WHITE)
    doc.paragraphs[-1].runs[0].font.small_caps = True
    for item in [
        "Recap of what we looked into since last meeting",
        "Does a state's default also punish its own city? (Hall's question)",
        "The canal-bond comparison — including one huge finding in Indiana",
        "A better source for our key date",
        "Two things we need your judgment on",
        "A proposed outline for the paper, if we're ready to start writing",
    ]:
        box_bullet(doc, item)
    end_box(doc, agenda_start, ACCENT_TINT, ACCENT, label_fill=ACCENT, label_text_white=True)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(10)

    say_this(
        doc,
        "Here's the tour: first a recap of what I dug into, then the city question you raised "
        "last time, then a big finding from the canal-bond comparison, a quick update on our key "
        "date, and then the two decisions I need — and if there's time, a proposed outline for "
        "the paper.",
    )

    # ------------------------------------------------------------------
    # What we investigated since last meeting
    # ------------------------------------------------------------------
    doc.add_heading("What We Investigated Since Last Meeting", level=1)
    doc.add_paragraph("Three loose ends from last time — here's what I found on each.")
    say_this(doc, "So let me start with three things that were still open after our last conversation.")

    doc.add_heading("Bank-held vs. traded bonds investigation", level=2)
    add_bullets(doc, [
        "You flagged that some banks were required to hold state bonds as backing for their own "
        "bank notes — meaning a bond could look \"untraded\" for reasons that have nothing to do "
        "with what investors thought of it. I checked every bond we use for unusually quiet "
        "trading.",
        "One bond (a New York canal bond) is a good match: New York had a real 1838 law requiring "
        "banks to buy and hold bonds like this one, and the timing lines up.",
        "One other quiet bond (an Ohio bond) looks like a different story — trading just moved to "
        "a more popular Ohio bond, not bank-holding.",
        "Two other quiet bonds are still a mystery — not enough evidence either way.",
        "Bottom line: doesn't change any of our results — none of the four are bonds we lean on "
        "for our main findings — but it's an honest caveat to have on record.",
    ])
    say_this(
        doc,
        "First, the bank-holding question you raised — whether some of our quiet bonds were just "
        "sitting in bank vaults instead of trading. Turns out it's a mixed bag: one New York bond "
        "really does look bank-held, one Ohio bond looks more like trading just moved elsewhere, "
        "and two others I honestly can't explain. None of it touches our main results, but I "
        "wanted you to know I checked.",
    )

    doc.add_heading("Alabama railroad bond flag", level=2)
    add_bullets(doc, [
        "A source mentioned Alabama paying bondholders only about 30 cents on the dollar on some "
        "railroad bonds — I wasn't sure if that was part of our 1840s story.",
        "It's not. That's a separate, later episode — Alabama's Reconstruction-era railroad bonds "
        "from the 1870s, thirty years after our data ends. No overlap with our project.",
    ])
    say_this(
        doc,
        "Quick one — I ran into a mention of Alabama paying bondholders 30 cents on the dollar, "
        "and wanted to make sure that wasn't part of our story. It's not — that's a totally "
        "separate railroad-bond episode from the 1870s. Good to rule out, doesn't touch anything.",
    )

    doc.add_heading("Ohio yield-rise claim", level=2)
    add_bullets(doc, [
        "You'd mentioned Ohio's interest rate rose in the 1840s, even though Ohio never "
        "defaulted — I checked whether that meant the \"safe\" state was getting punished too.",
        "It's real, but at the wrong moment. Ohio's rate spiked in early 1842 — that's the "
        "original 1837 crash working through the system, almost a year before the federal "
        "government's actual no-bailout announcement in 1843. By the time the announcement "
        "happened, Ohio's rate had already come back down to normal.",
    ])
    say_this(
        doc,
        "You'd mentioned Ohio's rates rising in the 1840s as maybe a sign the safe state got hit "
        "too. I checked — it's real, but it's the wrong moment. It's the 1837 crash itself, not "
        "the 1843 announcement: Ohio's rate had already spiked and come back down before the "
        "government even made its no-bailout statement. So it's not contagion from the policy, "
        "it's just the original panic.",
    )

    # ------------------------------------------------------------------
    # City contagion
    # ------------------------------------------------------------------
    doc.add_heading("Does a State's Default Punish Its Own City Too?", level=1)
    intro_p = doc.add_paragraph()
    intro_p.add_run(
        "This is Hall's question, in plain terms: if a state defaults, does the punishment spread "
        "to something nearby that didn't do anything wrong — specifically, the state's own city? "
    )
    note_run = intro_p.add_run(
        "(Note to self: the original notes said \"Pennsylvania vs. another state\" — that was a "
        "mix-up. The real question is Philadelphia, the city, versus Pennsylvania, the state it "
        "sits inside.)"
    )
    note_run.italic = True
    note_run.font.color.rgb = MUTED
    say_this(
        doc,
        "Next, the question you raised last time, Hall — does a state's default spread to its own "
        "city, even though the city didn't do anything wrong? I've got four pieces of evidence on "
        "this now.",
    )

    doc.add_heading("Philadelphia vs. Pennsylvania", level=2)
    add_bullets(doc, [
        "Before the announcement, Philadelphia's rate was already a bit lower than Pennsylvania's "
        "— about 1.6 percentage points lower.",
        "After the announcement, that gap widened to about 2.7 percentage points.",
        "Why: Pennsylvania's own rate shot up — toward 20% around its own default — while "
        "Philadelphia's rate barely moved, staying around 5%.",
    ])
    add_chart(
        doc, "chart_city_vs_state.png",
        "Philadelphia's rate (the city, purple) barely reacts to Pennsylvania's default "
        "(the state, blue).",
    )
    say_this(
        doc,
        "First data point: Philadelphia versus Pennsylvania. Before the announcement, the city "
        "was already borrowing a bit cheaper than the state — about a point and a half. After, "
        "that gap widened to about two and a half points, entirely because Pennsylvania's own "
        "rate spiked toward 20% while Philadelphia barely moved.",
    )

    doc.add_heading("Pittsburgh vs. Pennsylvania", level=2)
    add_bullets(doc, [
        "Same test, a second Pennsylvania city: Pittsburgh shows almost the identical gap, about "
        "2.65 percentage points. Same story, second confirmation.",
    ])
    add_chart(
        doc, "chart_pittsburgh_vs_pa.png",
        "Pittsburgh (the city) vs. Pennsylvania (the state) — same pattern as Philadelphia.",
    )
    say_this(
        doc,
        "I wanted to make sure Philadelphia wasn't a fluke, so I checked a second Pennsylvania "
        "city, Pittsburgh — same result, almost exactly the same gap. Two cities, same pattern.",
    )

    doc.add_heading("Cincinnati vs. Ohio — a control", level=2)
    add_bullets(doc, [
        "Then I checked a city in a state that never defaulted — Cincinnati inside Ohio. No gap "
        "at all, basically zero.",
        "That's actually a good sign: the gap only shows up when the state itself is genuinely in "
        "trouble. Ohio was fine, so Cincinnati was fine too.",
    ])
    add_chart(
        doc, "chart_cincinnati_vs_ohio.png",
        "Cincinnati (the city) vs. Ohio (the state) — no gap, because Ohio never defaulted.",
    )
    say_this(
        doc,
        "As a control, I checked Cincinnati inside Ohio — a state that never defaulted. No gap at "
        "all, basically flat. Which is reassuring: it means the city/state gap we're seeing isn't "
        "some generic pattern, it only shows up specifically where the state itself was in real "
        "trouble.",
    )

    doc.add_heading("Second Philadelphia bond", level=2)
    add_bullets(doc, [
        "One gap in the Philadelphia story: our main bond had a missing stretch of data right "
        "after the announcement, so we couldn't see the city's immediate reaction.",
        "Found a second Philadelphia bond that fills exactly that gap — same story, actually a "
        "bigger gap (about 6 percentage points), right in the window we couldn't see before.",
    ])
    say_this(
        doc,
        "One more thing on Philadelphia — our original bond had a hole in the data right after "
        "the announcement, so we couldn't see the city's immediate reaction. I found a second "
        "Philadelphia bond that covers exactly that missing window, and it shows the same "
        "no-punishment story — actually a bigger gap, about 6 points — right when it mattered "
        "most.",
    )

    doc.add_heading("Summary", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "Bottom line: the state gets punished; the city right next to it mostly doesn't. Three "
        "data points from Pennsylvania's own cities agree, and one clean comparison (Cincinnati) "
        "confirms the gap only shows up when the state is actually in trouble."
    )
    run.bold = True
    say_this(
        doc,
        "So overall: state gets punished, city doesn't. Three separate pieces of evidence from "
        "Philadelphia and Pittsburgh all agree, and Cincinnati as a control shows the gap only "
        "appears when the state itself is actually under stress. I think that's a pretty clean "
        "answer to your question.",
    )

    # ------------------------------------------------------------------
    # NEW (2026-09-02): Why didn't the city itself default?
    # ------------------------------------------------------------------
    doc.add_heading("Why Didn't the City Itself Default?", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "A different question from the one just above. That section showed the market didn't "
        "punish Philadelphia for Pennsylvania's default. This one asks why Philadelphia never got "
        "into trouble in the first place — why the city itself didn't default alongside the state. "
    )
    p.add_run(
        "Worth keeping these two separate: \"investors didn't punish the city\" is the price "
        "finding above; \"the city stayed solvent because it was a legally separate entity\" is "
        "this section."
    ).italic = True
    add_bullets(doc, [
        "Philadelphia was legally its own thing. A city in this era was a separate corporation "
        "from the state — its own budget, its own power to tax and borrow, not a branch of the "
        "state government.",
        "The state's debt paid for state projects. Pennsylvania borrowed to build the Main Line "
        "— a state-owned canal-and-railroad route from Philadelphia to Pittsburgh — plus other "
        "state canals and investments in banks. Philadelphia the city was never on the hook for "
        "any of that; it wasn't a co-signer or a guarantor.",
        "The city's debt paid for city things. Water works, gas works, wharves, public buildings "
        "— paid for out of the city's own property taxes and water and gas fees. That revenue "
        "was local and held up fine; it didn't depend on the state canal tolls and state taxes "
        "that collapsed.",
        "Different entity, different money — so the city never got pulled into the state's mess.",
    ])
    muted_note_bullet(
        doc,
        "Note to self: no single old source comes right out and says this in one sentence — it's "
        "pieced together from how the finances were structured plus our own price data. And I "
        "only checked the legal setup for Philadelphia specifically; Pittsburgh's price pattern "
        "matches but I didn't separately chase down its legal structure.",
    )
    say_this(
        doc,
        "One more thing worth adding here — why didn't the city default, not just why didn't the "
        "market punish it. Philadelphia was legally its own thing, separate from the state, with "
        "its own taxing power and its own budget. The state's debt paid for state canals and "
        "railroads that Philadelphia had nothing to do with; the city's debt paid for city "
        "things, funded by the city's own taxes. Different entity, different money, so it never "
        "actually got pulled into the state's mess in the first place. I'll flag this is "
        "inference from how the finances were structured, not a single source that spells it out "
        "directly — but it's a solid, consistent picture.",
    )

    # ------------------------------------------------------------------
    # Canal / robustness
    # ------------------------------------------------------------------
    doc.add_heading("The Canal-Bond Comparison — a Second Check on the Same Question", level=1)
    add_bullets(doc, [
        "Some state debt wasn't a plain loan to the state — it was backed by canal toll money "
        "specifically. We checked those separately, as a second, independent test.",
        "Indiana is the standout: in 1847 Indiana split its canal debt into two slices — one "
        "that got paid first (\"preferred\") and one that got paid only after the first slice was "
        "covered (\"deferred\"). The market priced that difference enormously: the first-in-line "
        "slice traded at 2-4 times the price of the second slice, a roughly 40-percentage-point "
        "interest-rate gap. That's the single biggest number in the whole project.",
        "Checked this wasn't a fluke — it holds up across 146 separate price readings over almost "
        "three years, not a handful of odd days.",
        "Two small data cleanups made along the way, flagged rather than hidden: one bond has a "
        "long gap in its price history (shown as a dotted line, not a straight line, so it "
        "doesn't look like a smooth climb that never happened); one single price reading looks "
        "like a typo in the original 1840s records and is excluded from the chart, though kept in "
        "the raw data.",
    ])
    add_chart(
        doc, "chart_canal_robustness.png",
        "Left: New York's canal bonds. Right: Indiana's first-in-line slice (green) vs. "
        "second-in-line slice (orange) of the same 1847 restructured loan.",
    )
    say_this(
        doc,
        "Now the canal-bond comparison — a second, independent test using a different kind of "
        "state debt, bonds backed specifically by canal toll money. The headline is Indiana: in "
        "1847 they split their canal debt into a slice that got paid first and a slice that got "
        "paid second, and the market priced that difference by about 40 percentage points — by "
        "far the biggest number in the whole project. I stress-tested it the same way I caught a "
        "problem in an earlier finding, and it holds up completely.",
    )

    # ------------------------------------------------------------------
    # NEW (2026-09-02): Indiana payment order -- Hall's follow-up question
    # ------------------------------------------------------------------
    doc.add_heading("The Indiana Payment Order — Answering Hall's Question", level=1)
    doc.add_paragraph(
        "Last meeting, Hall asked which of Indiana's canal bonds actually got paid first. Here's "
        "the clean version."
    )
    add_bullets(doc, [
        "The line-up: two bonds got paid first (the ones labeled \"preferred\"), and two got paid "
        "only after those (the \"deferred\" ones). A fifth bond, labeled just \"Indiana Canal,\" "
        "doesn't fit either group — it traded far higher than all four and looks like a separate, "
        "better-secured instrument, probably the part of the debt the state itself still stood "
        "behind.",
        "There's even a pecking order inside \"preferred\": the plain preferred bond traded higher "
        "than the one labeled \"special preferred.\" So the market wasn't just splitting these "
        "into two tiers — it was pricing finer distinctions than that.",
        "What we can't tell: whether a shortfall got split strictly (first group paid in full "
        "before the second group sees a cent) or proportionally (both groups take a haircut at "
        "the same time). No source we found states the rule.",
        "What the prices themselves suggest: the deferred bonds traded at a real, positive price "
        "— not near zero — so the market expected those holders to get something eventually. And "
        "even the preferred bonds only traded around 20 to 50 cents on the dollar, well below "
        "full value. So both groups were priced as taking a loss — not a clean \"the senior bond "
        "gets everything, the junior bond gets wiped out\" picture.",
        "When the bonds were issued: couldn't pin it down. The codebook has no issue-date field, "
        "and the legislative histories don't give a rollout schedule. The five bonds first show "
        "up on the exchange at staggered dates in 1850–51 — but that's three years after the 1847 "
        "restructuring act, and reflects when they started trading, not when they were issued. "
        "This is a genuine gap in the historical record, not something more digging is likely to "
        "fix.",
    ])
    say_this(
        doc,
        "So last time you asked which Indiana bonds got paid first — I've got a clear answer now. "
        "Two of them were \"preferred\" and got paid first; two were \"deferred\" and got paid "
        "after; and there's a fifth one that doesn't really fit either bucket and looks like a "
        "separate, safer instrument. There's even a sub-ranking inside the preferred group, so "
        "the market was pricing fine distinctions. What I can't tell you is the exact rule for "
        "how a shortfall gets split between them — nothing in the historical record spells that "
        "out — but the price data itself is telling: both the preferred and the deferred bonds "
        "traded well below full value, which suggests the market expected both sides to take some "
        "kind of hit, not a clean \"senior gets everything, junior gets nothing\" split. And I "
        "couldn't pin down exactly when the bonds were issued — the records just don't have it.",
    )

    # ------------------------------------------------------------------
    # Anchor date
    # ------------------------------------------------------------------
    doc.add_heading("A Better Source for Our Key Date", level=1)
    add_bullets(doc, [
        "We've been using February 11, 1843 as the date the federal government made clear it "
        "wouldn't bail states out. Until now that rested only on other historians' books, not an "
        "original document.",
        "I found and read an actual page from the original 1843 Congressional record — a real "
        "debate in Congress on exactly this topic, with the right people talking about the right "
        "issue, sitting right next to another date we'd already confirmed.",
    ])
    muted_note_bullet(
        doc,
        "Note to self: one honest caveat — the exact day inside that multi-day debate isn't 100% "
        "pinned down, but everything else lines up.",
    )
    say_this(
        doc,
        "Quick update on our key date, February 11, 1843. It used to rest only on other "
        "historians' accounts. I tracked down and actually read a page from the original 1843 "
        "Congressional record itself — a real debate on exactly this topic with the right people "
        "in the room. I can't swear to the exact day inside that multi-day debate, but everything "
        "else about it fits.",
    )

    # ------------------------------------------------------------------
    # Two decisions
    # ------------------------------------------------------------------
    doc.add_heading("Two Things We Need Your Judgment On", level=1)
    judgment_start = start_box(doc)
    box_heading(doc, "1. Alabama: default or not?", size=12, color=DECISION_TEXT)
    box_bullet(
        doc,
        "We've had Alabama filed under \"defaulted\" this whole time, but growing evidence says "
        "it actually never defaulted — it raised taxes and sold off state bank assets to keep "
        "paying instead. Three separate sources now back this up. We need your call before it "
        "goes into the paper.",
    )
    box_heading(doc, "2. Indiana's second-in-line bondholders — paid less, or paid late?",
                size=12, color=DECISION_TEXT)
    box_bullet(
        doc,
        "Our 40-point Indiana finding is solid on price, but we can't tell from the historical "
        "record whether the second-in-line slice actually stopped getting interest, or just got "
        "paid after the first slice, on a delay. Doesn't change the finding — just how precisely "
        "we describe it.",
    )
    end_box(doc, judgment_start, DECISION_FILL, DECISION_BORDER)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(
        "Note to self, not a decision: Cincinnati and the second Philadelphia bond are our newest "
        "findings and haven't been double-checked as hard as everything else yet."
    )
    run.italic = True
    run.font.color.rgb = MUTED
    say_this(
        doc,
        "So here's what I actually need from you two. First — do we feel good calling Alabama "
        "\"risky but survived\" instead of \"defaulted\"? I've now got three sources pointing "
        "that way, but I know you weren't sure last time, so I wanted your read before I bake it "
        "into the paper. Second — on the Indiana finding, do you know whether that second-in-line "
        "slice of debt actually stopped paying interest, or just got paid later? It doesn't "
        "change the number, just how precisely I can describe it. And one more thing, not a "
        "decision — Cincinnati and the second Philadelphia bond are our newest findings and "
        "haven't been stress-tested as hard as everything else yet.",
    )

    # ------------------------------------------------------------------
    # Proposed paper structure
    # ------------------------------------------------------------------
    doc.add_heading("If We're Ready to Start Writing — a Proposed Outline", level=1)
    doc.add_paragraph("Not locked in — just a starting point to react to.").italic = True
    outline = [
        "Introduction — the question, tied to the modern EU \"blue bonds\" debate.",
        "Background — the 1837 crash, the 1840s state defaults, the federal refusal to bail "
        "anyone out.",
        "Data and method — old bond prices turned into interest rates, our three state buckets, "
        "and the two different before/after moments we test separately (the 1837 crash itself vs. "
        "the 1843 no-bailout announcement).",
        "Main result — does the no-bailout promise show up in state interest rates: Pennsylvania, "
        "Ohio, Alabama, Indiana, New York.",
        "City result — does a state's default punish its own city too.",
        "Robustness — the canal-bond comparison and the Indiana finding.",
        "Discussion — the market drawing careful lines instead of panicking blindly, tied back to "
        "the EU parallel.",
        "Conclusion and limitations — the gaps and open questions we're upfront about.",
    ]
    for item in outline:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(3)
    say_this(
        doc,
        "If you think we're in good shape, here's roughly how I'm picturing the paper — let me "
        "know if this order makes sense or if you'd restructure it. Nothing here is locked in, "
        "this is just a starting point to react to.",
    )

    # ------------------------------------------------------------------
    # Closing
    # ------------------------------------------------------------------
    doc.add_heading("Wrapping Up", level=1)
    say_this(
        doc,
        "So to sum up where we landed: I need a yes-or-no on Alabama, and your best guess on the "
        "Indiana coupon question, and if the paper outline looks right, I'll start drafting. I'll "
        "write up any changes from today and send it your way — thanks again for the time, this "
        "was really helpful.",
    )

    add_page_numbers(doc)

    out_path = OUTPUT_DIR / "meeting_prep_v4.docx"
    doc.save(str(out_path))
    print(f"saved -> {out_path}")
    print("\nEmbedded charts:")
    for c in embedded_charts:
        print(f"  - {c}")
    if missing_charts:
        print("\nMISSING charts (placeholder inserted):")
        for c in missing_charts:
            print(f"  - {c}")


if __name__ == "__main__":
    main()
