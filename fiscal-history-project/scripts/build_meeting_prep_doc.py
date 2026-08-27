"""Build meeting_prep_final.docx -- live-presentation meeting prep document.

Restructured (2026-08-27) around an explicit agenda: recap, what changed since
last meeting, the city-contagion ("disease") answer, the canal/robustness
comparison, the anchor-date upgrade, two open advisor decisions, and a
proposed paper outline. All numbers are pulled from PROJECT_CONTEXT.md
(source of record) as of this build. Talking-doc style: tight bullets, not a
paper -- this is presented live, not read silently.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "output"

INK = RGBColor(0x0B, 0x0B, 0x0B)
MUTED = RGBColor(0x52, 0x51, 0x4E)
BOX_FILL = "FFF7DB"
BOX_BORDER = "B8860B"
AGENDA_FILL = "EEF2F7"
AGENDA_BORDER = "6B7A8F"

embedded_charts = []
missing_charts = []


def add_bullets(doc, items, bold_lead=None):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_chart(doc, filename, caption):
    path = OUTPUT_DIR / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.3))
        embedded_charts.append(str(path))
        add_caption(doc, caption)
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[chart not found — check output/: {filename}]")
        run.bold = True
        run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
        missing_charts.append(filename)


def add_toc(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose \"Update Field\" to generate the table of contents."
    fldChar_sep.append(placeholder)
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fldChar_begin)
    r.append(instrText)
    r.append(fldChar_sep)
    r.append(fldChar_end)


# --- Bordered/highlighted "box" implemented as bordered paragraphs, not a
# table. A single-cell table clipped its final line under Pages' docx
# renderer even with cantSplit set (confirmed the content was present in the
# XML via python-docx -- a Pages table-cell rendering bug, not a data bug).
# Paragraph borders (w:pBdr) don't have that failure mode in either Word or
# Pages, so the box is built by bordering a contiguous run of paragraphs.


def start_box(doc):
    return len(doc.paragraphs)


def end_box(doc, start_index, fill_hex, border_hex, indent_pt=14):
    # All box paragraphs use the "Normal" style with identical explicit
    # indent/spacing (not "List Bullet", whose own style-level indent and
    # spacing don't line up with plain paragraphs and broke the border's
    # visual continuity into disconnected segments).
    paras = doc.paragraphs[start_index:]
    n = len(paras)
    for i, p in enumerate(paras):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        edges = ["left", "right"]
        if i == 0:
            edges.append("top")
        if i == n - 1:
            edges.append("bottom")
        for edge in edges:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "18")
            el.set(qn("w:space"), "8")
            el.set(qn("w:color"), border_hex)
            pBdr.append(el)
        pPr.append(pBdr)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        pPr.append(shd)
        p.paragraph_format.left_indent = Pt(indent_pt)
        p.paragraph_format.right_indent = Pt(indent_pt)
        p.paragraph_format.space_before = Pt(8 if i == 0 else 4)
        p.paragraph_format.space_after = Pt(8 if i == n - 1 else 4)


def box_heading(doc, text, first=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return p


def box_bullet(doc, text):
    p = doc.add_paragraph()
    p.add_run("•  " + text)
    return p


def box_paragraph(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    return p


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = INK

    # ------------------------------------------------------------------
    # 1. Title + recap
    # ------------------------------------------------------------------
    doc.add_heading("1840s State Debt Credibility — Progress Since Last Meeting", level=0)

    add_toc(doc)

    doc.add_paragraph(
        "Research question: did bond markets treat the federal government's refusal to bail out "
        "defaulting states in the early 1840s as a credible commitment? The test is whether states "
        "that actually defaulted show persistently higher bond yields after the federal no-bailout "
        "signal than states that never defaulted, with states that were genuinely at risk but "
        "ultimately survived landing in between. It's been about a month since we last met — this "
        "document is meant to stand on its own as a full walkthrough of everything done since then, "
        "not to assume we're picking up where we left off."
    )
    doc.add_page_break()

    # ------------------------------------------------------------------
    # 2. Agenda box
    # ------------------------------------------------------------------
    doc.add_heading("Agenda", level=1)
    agenda_start = start_box(doc)
    box_heading(doc, "Today's flow", first=True)
    for item in [
        "Recap of what was investigated since last meeting",
        "City-contagion findings (Hall's \"disease\" question)",
        "Canal/robustness comparison + Indiana tranche finding",
        "Anchor-date upgrade",
        "Two open questions needing a decision",
        "Proposed structure for the paper, if we're ready to move to writing",
    ]:
        box_bullet(doc, item)
    end_box(doc, agenda_start, AGENDA_FILL, AGENDA_BORDER)
    doc.add_paragraph()

    # ------------------------------------------------------------------
    # 3. What we investigated since last meeting
    # ------------------------------------------------------------------
    doc.add_heading("What We Investigated Since Last Meeting", level=1)

    doc.add_heading("Bank-held vs. traded bonds investigation", level=2)
    add_bullets(doc, [
        "You flagged that state-chartered banks were sometimes required to hold state bonds as "
        "backing for their own bank notes — a bond held this way would show sparse market data for "
        "reasons that have nothing to do with investor sentiment. Built a trade-density flag across "
        "all 23 primary/secondary bond codes and checked the two already-known sparse bonds plus "
        "the two new ones the flag turned up.",
        "S-1750 (New York canal, 0.36 obs/month, a 1,666-day gap) is the best circumstantial fit: "
        "New York's 1838 Free Banking Act required banks to buy and deposit state bonds against "
        "note issuance, and the timing lines up (gap starts 1843, five years after the Act).",
        "S-2010 (Ohio, 0.60 obs/month) leans toward a different explanation — liquidity concentrating "
        "onto a single benchmark issue (S-2080) rather than bank absorption. Ohio's comparable law "
        "(the 1845 Kelley Act) postdates the start of this gap by about two years, a timing mismatch.",
        "Two remain unexplained: S-0040 (Alabama, no maturity data in the codebook at all, only a "
        "weak circumstantial link) and S-2190 (Ohio canal, just 2 observations, both from 1825 — a "
        "data-coverage limit before the study window even opens, not a trading-pattern finding).",
        "Bottom line: mixed, honest result — not all four sparse bonds share one cause. Doesn't "
        "change any existing chart (none of the four are used as a state's sole evidence anywhere), "
        "but it's a real caveat on data quality for the still-secondary canal comparison.",
    ])

    doc.add_heading("Alabama railroad bond flag", level=2)
    add_bullets(doc, [
        "A source we ran into mentioned Alabama paying bondholders roughly 30 cents on the dollar on "
        "railroad-guaranteed bonds — unclear at the time whether that was the same 1840s episode.",
        "Resolved: it's a distinct, later episode. Alabama's 1867 internal-improvements act had the "
        "governor endorse railroad bonds; the Alabama & Chattanooga Railroad defaulted in Jan 1871; "
        "Governor Houston's 1875-76 debt commission recognized $18M of claims, later cut to $12.5M, "
        "against original face claims around $30M — consistent with the \"~30 cents on the dollar\" "
        "figure, but for Reconstruction-era railroad-aid bonds, not our 1840s state bonds.",
        "No possible overlap with our data: Alabama's price series (S-0030/S-0040) runs 1842-1853; "
        "the railroad bonds weren't issued until 1867, 14 years after our data ends. Confirmed from "
        "the same Wallis NBER paper already in our reading list, which separates the two episodes in "
        "two distinct passages. Does not touch or complicate the Alabama reclassification question.",
    ])

    doc.add_heading("Ohio yield-rise claim", level=2)
    add_bullets(doc, [
        "You'd mentioned Ohio (never defaulted, our \"safe\" bucket) saw its yield rise during the "
        "1840s — flagged as possible evidence contagion reached even the safe bucket. Checked at "
        "monthly resolution across all four Ohio primary bonds (S-2100/S-2110/S-2080/S-2010).",
        "Resolved: real move, wrong window. All four bonds show a synchronized price collapse from "
        "~90 (Oct 1841) to ~48-52, bottoming Mar 12-Apr 9 1842 — peak yields S-2100 16.02% "
        "(Mar 26 1842), S-2010 15.85% (Apr 9 1842), S-2110 12.83% (Mar 12 1842), S-2080 11.48% "
        "(Mar 12 1842).",
        "This is the 1837 panic window, not the 1843 policy window: the spike bottoms about 4.5 "
        "months before PA's actual default and about 11 months before the Feb 11 1843 no-bailout "
        "signal. Yields normalize back to a 6-7% baseline starting May 1843 — the opposite of what "
        "policy-driven contagion into the safe bucket would predict.",
        "Reads as a conflation of the panic window and the policy window — exactly the distinction "
        "this project has guarded against from the start. Useful confirmation that the panic/policy "
        "split is doing real analytical work, not just bookkeeping.",
    ])

    # ------------------------------------------------------------------
    # 4. City contagion — the "disease" question
    # ------------------------------------------------------------------
    doc.add_heading("City Contagion — The \"Disease\" Question", level=1)
    doc.add_paragraph(
        "Hall's framing: does state default spread like a disease to the surrounding area — "
        "specifically, does it spread downward to the defaulting state's own city? (One "
        "clarification: the original notes said \"Pennsylvania,\" but the actual question is "
        "Philadelphia — PA's own city — versus Pennsylvania the state, not Pennsylvania versus "
        "some other state. That was a note-taking slip we caught and corrected.)"
    )

    doc.add_heading("Philadelphia vs. Pennsylvania (C-1100)", level=2)
    add_bullets(doc, [
        "Pre-signal: 1.63pp gap (Pennsylvania 6.69%, Philadelphia 5.06%).",
        "Post-gap-overlap window (the only period both series actually have data, apples-to-apples): "
        "2.68pp gap (Pennsylvania 7.56%, Philadelphia 4.88%).",
        "The widening is driven entirely by Pennsylvania's own yield rising — Philadelphia's yield "
        "barely moves at all. The state spikes toward ~20% yield around its own default; "
        "Philadelphia stays flat around 5%.",
    ])
    add_chart(
        doc, "chart_city_vs_state.png",
        "Philadelphia (city) vs. Pennsylvania (state) yields, 1839-1844.",
    )

    doc.add_heading("Pittsburgh vs. Pennsylvania", level=2)
    add_bullets(doc, [
        "A second Pennsylvania city, checked to see if Philadelphia's result replicates: it does, "
        "almost exactly — 2.65pp gap (Pennsylvania 8.76%, Pittsburgh 6.11%), essentially identical "
        "in size and direction to Philadelphia's own 2.68pp.",
    ])
    add_chart(
        doc, "chart_pittsburgh_vs_pa.png",
        "Pittsburgh (city) vs. Pennsylvania (state) yields.",
    )

    doc.add_heading("Cincinnati vs. Ohio — negative control", level=2)
    add_bullets(doc, [
        "Cincinnati vs. Ohio comes in near-parity: -0.11pp (Ohio 6.25%, Cincinnati 6.36%) — "
        "essentially flat.",
        "This is NOT a third confirmation of no-spillover — Ohio never defaulted, so there's no "
        "state distress for Cincinnati to be insulated from in the first place. It's a baseline "
        "contrast: when the state's own credit is fine, its city tracks it with no gap at all. The "
        "~2.6-2.7pp gap only opens up for the two Pennsylvania cities, where the state itself was "
        "genuinely under stress.",
    ])
    add_chart(
        doc, "chart_cincinnati_vs_ohio.png",
        "Cincinnati (city) vs. Ohio (state) yields — negative control (Ohio never defaulted).",
    )

    doc.add_heading("Second Philadelphia bond (C-1260)", level=2)
    add_bullets(doc, [
        "C-1100 has a 343-day data gap (Feb 1843-Feb 1844) that swallows almost the entire immediate "
        "post-signal reaction period. A second Philadelphia city bond, C-1260, has real data exactly "
        "where C-1100 doesn't.",
        "Two figures reported together, since they trade off precision against coverage: 5.95pp "
        "(Philadelphia 5.59% vs. Pennsylvania 11.53%, n=19, Apr-Aug 1843) using the project's "
        "standard post-signal window, or 6.28pp (Philadelphia 5.63% vs. Pennsylvania 11.91%, n=24) "
        "using the full window that fills the gap (Feb 25-Aug 5 1843).",
        "This confirms the state/city gap holds even during the immediate post-signal reaction "
        "window, not just before and after it — closing the biggest caveat on the flagship "
        "Philadelphia finding.",
    ])

    doc.add_heading("Summary", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "The state gets punished; the city right next to it mostly doesn't — three consistent "
        "Pennsylvania-city data points (Philadelphia's C-1100, Philadelphia's C-1260, and "
        "Pittsburgh) and one clean negative control (Cincinnati, where the state itself was never "
        "under stress)."
    )
    run.bold = True

    # ------------------------------------------------------------------
    # 5. Canal / robustness comparison
    # ------------------------------------------------------------------
    doc.add_heading("Canal / Robustness Comparison", level=1)
    doc.add_paragraph(
        "A secondary comparison using revenue-pledged canal bonds instead of general-obligation "
        "bonds: New York's canal series (S-1750/S-1820/S-1950) and Indiana's 1847 Butler Bill "
        "restructuring tranches. Ohio's only canal bond (S-2190) is dropped — 2 observations, both "
        "from 1825, no data anywhere in the 1840s-50s study window."
    )
    add_bullets(doc, [
        "Headline number: Indiana's 1847 restructuring split its canal debt into a senior "
        "\"preferred\" tranche and a junior \"deferred\" tranche. The market priced that seniority "
        "split by a 40.04pp current-yield gap (preferred 17.96% vs. deferred 57.99%) — the largest "
        "spread in the whole project, confirmed directly in raw prices (preferred trades 2-4x "
        "deferred throughout), sustained across 146 observations, not a thin-sample artifact.",
        "Two data-integrity catches made along the way, flagged rather than silently patched: "
        "S-1750's known 1,666-day gap is rendered as a dotted, non-connecting chart segment, not a "
        "straight line implying a climb that never happened; an isolated S-1820 price of 160.00 "
        "(genuine in the raw source, almost certainly a transcription error) is excluded from the "
        "chart line but kept as-is in the underlying CSV.",
    ])
    add_chart(
        doc, "chart_canal_robustness.png",
        "Canal/robustness comparison: NY canal bonds and Indiana preferred vs. deferred tranches.",
    )

    # ------------------------------------------------------------------
    # 6. Anchor date upgrade
    # ------------------------------------------------------------------
    doc.add_heading("Anchor Date Upgrade", level=1)
    add_bullets(doc, [
        "The Feb 11 1843 no-bailout signal date now rests on a real primary source, not just "
        "secondary accounts: UNT's digitized Congressional Globe, Vol. 12, 27th Congress, 3rd "
        "session, pages 283-284.",
        "Content: a genuine House floor debate — \"Mississippi State Bonds\" — with Gwin (MS), "
        "Granger (Whig-NY), and Thompson (MS) disputing state bond authorization and federal "
        "responsibility for recognizing the debt.",
        "The exact day within the multi-day debate isn't pinned — neither page carries a visible "
        "date stamp — but content, participants, and page position (8-9 pages before an "
        "already-confirmed Feb 16 1843 page) all line up tightly.",
        "Upgrade: from \"two secondary sources only\" to \"secondary sources plus a real, on-topic "
        "primary-source read of the right participants and subject in the right tight date window.\"",
    ])

    # ------------------------------------------------------------------
    # 7. Two things we need your judgment on
    # ------------------------------------------------------------------
    doc.add_heading("Two Things We Need Your Judgment On", level=1)
    judgment_start = start_box(doc)

    box_heading(doc, "1. Alabama reclassification (Defaulted → Risky but survived)", first=True)
    box_paragraph(doc, "What we know:", bold=True)
    box_bullet(doc, "Three independent sources now support it: Wallis's NBER paper explicitly "
                     "excludes Alabama from the defaulters; two independent lists of 1840s "
                     "defaulting states both exclude Alabama.")
    box_paragraph(doc, "What we need from you:", bold=True)
    box_bullet(doc, "A yes/no — or explicit permission to treat this as open in the paper. Not "
                     "personally confirmed by either of you yet.")

    box_heading(doc, "2. Indiana's deferred tranche coupon status")
    box_paragraph(doc, "What we know:", bold=True)
    box_bullet(doc, "Could not determine from available literature whether the deferred "
                     "tranche stopped paying coupon entirely or just sat lower in priority. "
                     "Doesn't change the 40.04pp price-based finding — only how precisely it "
                     "should be worded.")
    box_paragraph(doc, "What we need from you:", bold=True)
    box_bullet(doc, "Any guidance from the literature — or confirmation it's fine to note as "
                     "unresolved in the paper.")

    end_box(doc, judgment_start, BOX_FILL, BOX_BORDER)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "Also worth flagging, not requiring approval: Cincinnati and the second Philadelphia bond "
        "(C-1260) are the newest, least-reviewed pieces of the analysis."
    )
    run.italic = True

    # ------------------------------------------------------------------
    # 8. Proposed paper structure
    # ------------------------------------------------------------------
    doc.add_heading("Proposed Paper Structure", level=1)
    doc.add_paragraph(
        "If we're ready to start writing — proposed here to confirm or adjust live, not a locked plan."
    ).italic = True
    outline = [
        "Introduction — the research question, tied to the Lustig \"Blue Bonds for Europe\" framing.",
        "Historical background — 1837 panic, 1840s state defaults, federal refusal to assume debt.",
        "Data and methodology — EH.net securities prices, three-bucket classification, panic-window "
        "vs. policy-window distinction, active-default override rule.",
        "Main result — does the no-bailout signal show up in state-level yields: PA/Ohio/Alabama/"
        "Indiana/NY comparison.",
        "City contagion result — does default punish the defaulting state's own city "
        "(Philadelphia/Pittsburgh/Cincinnati).",
        "Robustness — canal bonds, Indiana tranche split, reproducibility checks.",
        "Discussion — markets doing real credit analysis on separate legal/fiscal entities rather "
        "than generic panic; tie back to the EU \"blue bonds\" parallel.",
        "Conclusion, limitations — NYC/Brooklyn data insufficiency, bond seniority confound, "
        "single-day anchor-date uncertainty.",
    ]
    for i, item in enumerate(outline, start=1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)

    out_path = OUTPUT_DIR / "meeting_prep_final.docx"
    doc.save(str(out_path))
    print(f"saved -> {out_path}")
    print("\nEmbedded charts:")
    for c in embedded_charts:
        print(f"  - {c}")
    if missing_charts:
        print("\nMISSING charts (placeholder inserted):")
        for c in missing_charts:
            print(f"  - {c}")


if __name__ == "__main__":
    main()
